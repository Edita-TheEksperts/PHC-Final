"""
Builds PHC_Test_Dokumentation.docx — a human-tone, step-by-step QA document
covering all 48 feedback items from the 12.05.2026 audit. Each item has:
  • Was war das Problem?
  • Was haben wir geändert?
  • Wo und wie testen? (numbered steps)
  • Screenshot-Platzhalter mit Beschriftung
  • Erwartetes Ergebnis

Run:  python build_test_doc.py
Output: PHC_Test_Dokumentation.docx in the working directory.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SCREENSHOT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "screenshots")


def find_screenshots(num):
    """Return list of image paths matching this F-XX identifier.
    Accepts F-XX as-is or combined like 'F-13 / F-18' → checks F-13.png and F-18.png."""
    candidates = []
    for token in num.replace(" ", "").split("/"):
        for ext in ("png", "jpg", "jpeg"):
            p = os.path.join(SCREENSHOT_DIR, f"{token}.{ext}")
            if os.path.exists(p):
                candidates.append(p)
                break
    # dedupe while preserving order
    seen = set()
    out = []
    for c in candidates:
        if c not in seen:
            out.append(c)
            seen.add(c)
    return out


# ── Helpers ──────────────────────────────────────────────────────────────

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B99B5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def screenshot_placeholder(doc, caption):
    """Insert a styled placeholder box where the user will paste a screenshot."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_background(cell, "FAFCFF")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ Screenshot hier einfügen ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xB9, 0x9B, 0x5F)
    run.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(caption)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "dashed")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "B99B5F")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    doc.add_paragraph()  # spacer


def add_steps(doc, steps):
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(s)


def add_item(doc, num, title, prio, problem, change, where_to_test, expected, screenshot_caption):
    # Heading
    h = doc.add_heading(f"{num} — {title}", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x04, 0x43, 0x6F)

    # Priority badge line
    p = doc.add_paragraph()
    r = p.add_run(f"Priorität: {prio}     Status: Erledigt")
    r.bold = True
    if prio == "KRITISCH":
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    elif prio == "HOCH":
        r.font.color.rgb = RGBColor(0xD9, 0x82, 0x2B)
    else:
        r.font.color.rgb = RGBColor(0x04, 0x43, 0x6F)
    r.font.size = Pt(10)

    # Problem
    doc.add_heading("Was war das Problem?", level=3)
    doc.add_paragraph(problem)

    # Change
    doc.add_heading("Was wurde geändert?", level=3)
    doc.add_paragraph(change)

    # Where to test
    doc.add_heading("Wo und wie testen?", level=3)
    add_steps(doc, where_to_test)

    # Screenshot — embed if Playwright captured one, otherwise placeholder
    doc.add_heading("Screenshot", level=3)
    shots = find_screenshots(num)
    if shots:
        for path in shots:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Inches(6.0))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(f"Automatisch aufgenommen: {os.path.basename(path)}")
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    else:
        screenshot_placeholder(doc, screenshot_caption)

    # Expected
    doc.add_heading("Erwartetes Ergebnis", level=3)
    doc.add_paragraph(expected)

    # Divider
    sep = doc.add_paragraph()
    add_horizontal_line(sep)


# ── Document setup ───────────────────────────────────────────────────────

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


# ── Title page ───────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("PHC SYSTEM")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x04, 0x43, 0x6F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Update- & Test-Dokumentation")
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0xB9, 0x9B, 0x5F)

doc.add_paragraph()

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = intro.add_run("Behebung der 48 Feedback-Punkte aus dem Audit vom 12.05.2026")
r.font.size = Pt(12)
r.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Prime Home Care AG\n").bold = True
meta.add_run("Stand: 20.05.2026\n")
meta.add_run("Erstellt von: Fisnik Salihu\n")
meta.add_run("Adressat: Bettina, Silvain (Abnahme)")

doc.add_paragraph()
doc.add_paragraph()

# Vorwort
doc.add_heading("Vorwort", level=1)
p = doc.add_paragraph()
p.add_run(
    "Dieses Dokument beschreibt jeden der 48 Feedback-Punkte einzeln — was wir geändert haben, "
    "wo ihr es im laufenden System nachprüfen könnt und was als Ergebnis erwartet wird. "
    "Die Reihenfolge folgt dem Audit-PDF vom 12.05.2026, ist also nach Prozessablauf gruppiert "
    "(Buchung → Einsatz-Generierung → E-Mail → Kunden-Dashboard → Mitarbeiter-Dashboard → "
    "Admin-Dashboard → Allgemeines)."
)
p = doc.add_paragraph()
p.add_run(
    "Jeder Block hat denselben Aufbau: eine kurze Erklärung was vorher schief lief, eine "
    "knappe Beschreibung was wir verändert haben, eine Schritt-für-Schritt-Anleitung zum "
    "Nachvollziehen im Browser, ein Platzhalter wo ihr beim Testen einen Screenshot einfügt, "
    "und eine Zeile mit dem erwarteten Ergebnis. So lässt sich die Datei direkt als "
    "Abnahmeprotokoll verwenden."
)
p = doc.add_paragraph()
p.add_run(
    "Test-Umgebung: bitte die Staging-URL verwenden (nicht die Live-Domain). Login-Daten für "
    "die drei Rollen (Admin, Kunde, Mitarbeiter/Kandidat) liegen separat im PHC-Passwortordner. "
    "Sollte ein Schritt nicht funktionieren wie beschrieben — bitte nicht annehmen, sondern "
    "im Sheet markieren und mich (Fisnik) direkt benachrichtigen."
)

doc.add_page_break()


# ── Inhaltsverzeichnis ───────────────────────────────────────────────────

doc.add_heading("Übersicht", level=1)

overview = [
    ("Phase 1", "Kritische Bugs", "F-09, F-12, F-13/18, F-06, F-24, F-25, F-38, F-48/23/42"),
    ("Phase 2", "E-Mail-Infrastruktur", "F-14, F-15, F-16, F-17, F-43"),
    ("Phase 3", "Datenintegrität & Anzeige", "F-01, F-02, F-03/28, F-10, F-11, F-26, F-34, F-35, F-39"),
    ("Phase 4", "UI/UX Umbau", "F-20, F-22, F-27, F-29, F-31, F-32, F-33, F-36, F-40, F-42, F-44, F-45"),
    ("Phase 5", "Feinschliff", "F-04, F-05, F-07, F-08, F-19, F-21, F-30, F-37, F-41, F-46, F-47"),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "Phase"
hdr[1].text = "Thema"
hdr[2].text = "Punkte"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
for ph, th, pts in overview:
    row = tbl.add_row().cells
    row[0].text = ph
    row[1].text = th
    row[2].text = pts

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run(
    "Total: 48 Punkte. 45 vollständig implementiert, 3 als Geschäfts-/QA-Aufgabe "
    "(F-15 finale E-Mail-Texte, F-19 manueller Smoke-Test, F-04 visuelle Stichprobe)."
).italic = True

doc.add_page_break()


# ── Phase 1: Kritische Bugs ──────────────────────────────────────────────

doc.add_heading("Phase 1 — Kritische Bugs", level=1)
doc.add_paragraph(
    "Die acht Punkte in dieser Phase waren der Grund warum wir den Go-Live ursprünglich "
    "verschoben haben. Sie betreffen Sicherheit, Datenschutz, Datenintegrität und finanzielle "
    "Prozesse. Alle sind jetzt erledigt und durch Code-Reviews bestätigt."
)
doc.add_paragraph()

add_item(
    doc, "F-12", "E-Mails gehen jetzt an den Auftraggeber", "KRITISCH",
    problem=(
        "Bis vor kurzem schickte das System sämtliche Buchungs- und Termin-E-Mails an die "
        "Mailadresse der betreuten Person. Das ist nicht nur falsch — die betreute Person ist "
        "in den meisten Fällen gar nicht der/die Auftraggeber:in und hat selten überhaupt eine "
        "E-Mail-Adresse — sondern auch ein klarer Verstoss gegen das Schweizer DSG."
    ),
    change=(
        "In src/lib/recipientEmail.js wurde ein zentraler Helfer eingeführt, der bei jeder "
        "kundenrelevanten E-Mail zuerst die requestEmail (Auftraggeber) nimmt und nur dann auf "
        "User.email zurückfällt, wenn die Auftraggeber-Adresse leer ist. Die mailer.js und alle "
        "API-Endpunkte (Terminbestätigung, Stornierung, Kündigung etc.) wurden auf diesen "
        "Helfer umgestellt."
    ),
    where_to_test=[
        "Loggen Sie sich als Admin im Staging-System ein.",
        "Erstellen Sie testweise einen neuen Kunden, bei dem die betreute Person eine andere "
        "E-Mail-Adresse hat als der/die anfragende Person (Auftraggeber).",
        "Schliessen Sie den Buchungsprozess ab.",
        "Prüfen Sie im Mailserver oder im jeweiligen Postfach, an welche Adresse die "
        "Willkommens-E-Mail gegangen ist.",
        "Wiederholen Sie den Test mit Terminbestätigung und Stornierung.",
    ],
    expected=(
        "Alle System-E-Mails landen in der Auftraggeber-Mailbox. Die betreute Person erhält "
        "keine E-Mails (es sei denn beide Adressen sind absichtlich identisch)."
    ),
    screenshot_caption=(
        "Empfänger-Adresse im Postfach des Auftraggebers (z. B. Gmail-Detailansicht mit "
        "sichtbarem 'An:'-Feld)."
    ),
)

add_item(
    doc, "F-38", "Session-Bug: kein 'fremder User' mehr nach Token-Ablauf", "KRITISCH",
    problem=(
        "Bettina hat berichtet, dass im Admin-Dashboard plötzlich Daten eines anderen Users "
        "auftauchten (Bruno Roth, Sandra Keller). Ursache: ein abgelaufener JWT wurde nicht "
        "konsequent ausgewertet, und beim Logout blieb localStorage teilweise gefüllt — "
        "dadurch hat der Browser beim nächsten Request mit alten Daten gerendert."
    ),
    change=(
        "Die Middleware (src/middleware.js) prüft jetzt aktiv das exp-Feld des JWT, "
        "src/lib/clientAuth.js stellt isTokenExpired() und clearAuthStorage() zur Verfügung. "
        "Jede authentifizierte Anfrage läuft über authFetch() — bei einer 401-Antwort wird "
        "localStorage komplett geleert und auf /login umgeleitet."
    ),
    where_to_test=[
        "Loggen Sie sich als Admin ein.",
        "Lassen Sie den Tab mindestens so lange offen, bis das Token abgelaufen ist "
        "(Token-Lebensdauer steht in .env, üblicherweise 8 Stunden — alternativ in den "
        "DevTools 'userToken' aus localStorage entfernen).",
        "Klicken Sie auf einen beliebigen Menüpunkt im Admin-Dashboard.",
        "Loggen Sie sich anschliessend bewusst aus (Logout-Button).",
        "Öffnen Sie die DevTools → Application → Local Storage und prüfen Sie ob die "
        "Schlüssel userToken, email, role wirklich entfernt sind.",
    ],
    expected=(
        "Bei abgelaufenem Token erfolgt sofort ein Redirect auf /login. Nach Logout ist "
        "localStorage leer. Es werden nie mehr Daten eines anderen Users gezeigt."
    ),
    screenshot_caption=(
        "DevTools-Ansicht des Local Storage nach erfolgtem Logout — alle PHC-Schlüssel sind weg."
    ),
)

add_item(
    doc, "F-09", "Einsätze finden an den richtigen Wochentagen statt", "KRITISCH",
    problem=(
        "Wenn ein Kunde Mo/Mi/Fr gebucht hatte, generierte das System alle Folge-Einsätze auf "
        "denselben Wochentag. Grund: beim Setzen des Beginndatums wurde der Wochentag der "
        "ersten Schedule-Zeile automatisch überschrieben — danach hatte das System nur noch "
        "diesen einen Tag in der Liste und sortierte alles drauf."
    ),
    change=(
        "In src/pages/register-client.js (Bereich um Zeile 1990) wurde die Logik so umgebaut, "
        "dass schedules[0].day nur dann mit dem Wochentag des Beginndatums vorbelegt wird, "
        "wenn dieses Feld leer ist und der Wochentag noch nicht in einer anderen Zeile steht. "
        "Existierende Auswahlen werden nie mehr stillschweigend überschrieben."
    ),
    where_to_test=[
        "Öffnen Sie als (anonymer) Besucher die Buchungsstrecke (z. B. /betreuung-zuhause-organisieren).",
        "Wählen Sie als Frequenz 'wöchentlich' und tragen Sie drei Schedule-Zeilen ein: "
        "Montag, Mittwoch, Freitag.",
        "Wählen Sie ein Beginndatum, das zufällig auf einen anderen Wochentag fällt (z. B. einen Dienstag).",
        "Schliessen Sie die Buchung ab und loggen Sie sich anschliessend als dieser Kunde ein.",
        "Im Kunden-Dashboard unter 'Meine Verträge' (oder 'Termine') die ersten 4–6 Einsätze prüfen.",
    ],
    expected=(
        "Die generierten Termine fallen abwechselnd auf Mo, Mi, Fr — nicht alle auf denselben Tag. "
        "Das Beginndatum hat den ersten Schedule-Eintrag nicht überschrieben."
    ),
    screenshot_caption=(
        "Termin-Liste im Kunden-Dashboard mit sichtbaren unterschiedlichen Wochentagen."
    ),
)

add_item(
    doc, "F-13 / F-18", "Mitarbeiter-Zuweisungs-E-Mail kommt nur noch einmal", "KRITISCH",
    problem=(
        "Bei jeder Zuweisung wurde die E-Mail an den Mitarbeiter zweifach versendet. Ursache "
        "war ein doppelter API-Call (Doppelklick / React-Replay), der zwei Assignment-Zeilen "
        "anlegte und je eine E-Mail auslöste."
    ),
    change=(
        "Im Endpoint src/pages/api/admin/assign-employee.js gibt es jetzt einen Idempotenz-"
        "Guard: vor dem Anlegen wird per findFirst geprüft, ob eine Assignment-Zeile mit "
        "derselben Kombination (userId, employeeId, scheduleId) bereits existiert. Falls ja, "
        "kehrt der Endpoint mit { deduplicated: true } zurück — ohne zweite E-Mail."
    ),
    where_to_test=[
        "Loggen Sie sich als Admin ein.",
        "Öffnen Sie die Einsätze-Übersicht und wählen Sie eine Buchung ohne zugewiesenen Mitarbeiter.",
        "Klicken Sie auf 'Mitarbeiter zuweisen' und wählen Sie einen Kandidaten aus.",
        "Klicken Sie absichtlich zweimal hintereinander schnell auf 'Speichern' / 'Zuweisen'.",
        "Prüfen Sie das Postfach des zugewiesenen Mitarbeiters.",
    ],
    expected=(
        "Genau eine 'Neuer Einsatz – Bitte bestätigen'-E-Mail im Postfach des Mitarbeiters. "
        "Im Admin-Dashboard ist die Assignment-Zeile nur einmal sichtbar."
    ),
    screenshot_caption=(
        "Mitarbeiter-Postfach mit nur einer Zuweisungs-E-Mail (Eingang gefiltert auf Betreff "
        "'Neuer Einsatz')."
    ),
)

add_item(
    doc, "F-06", "Bewerbungsprozess endet jetzt auf einer Bestätigungsseite", "KRITISCH",
    problem=(
        "Nach Abschluss der Bewerbung landete der Bewerber wieder auf der letzten Formularseite, "
        "und die Bewerbungs-E-Mail kam nicht an. Es entstand der Eindruck einer Endlosschleife."
    ),
    change=(
        "Eine neue Seite src/pages/bewerbung-erfolgreich.js wurde angelegt. Nach erfolgreichem "
        "Submit wird der FormContext-State zurückgesetzt und auf diese Seite weitergeleitet. "
        "Die Bestätigungs-E-Mail (Template 'applicantConfirmation') wird im fire-and-forget-"
        "Modus nach dem DB-Save versendet, also nicht mehr durch UI-Probleme blockiert."
    ),
    where_to_test=[
        "Öffnen Sie die Bewerbungsstrecke /employee-register (anonym).",
        "Füllen Sie alle Schritte vollständig aus und reichen Sie die Bewerbung ein.",
        "Beobachten Sie, auf welche Seite umgeleitet wird.",
        "Prüfen Sie das Bewerber-Postfach auf die Bestätigungs-E-Mail (Betreff: "
        "'Ihre Bewerbung bei Prime Home Care AG').",
    ],
    expected=(
        "Weiterleitung auf /bewerbung-erfolgreich mit einer freundlichen Bestätigung. "
        "E-Mail liegt innerhalb von ca. 1 Minute im Postfach."
    ),
    screenshot_caption=(
        "Bestätigungsseite /bewerbung-erfolgreich im Browser, danach im selben Screenshot oder "
        "darunter die zugehörige E-Mail im Postfach."
    ),
)

add_item(
    doc, "F-48 / F-23 / F-42", "Speichern funktioniert in allen Bereichen", "KRITISCH",
    problem=(
        "Termine konnten bearbeitet, aber nicht gespeichert werden. Auch im Admin-Bereich "
        "(Einsätze) führte 'Speichern' ins Leere. Es gab keine Fehlermeldungen — der Klick "
        "wirkte einfach folgenlos."
    ),
    change=(
        "Die PUT-Route src/pages/api/appointments.js wurde fertiggestellt (komplette "
        "Update-Logik inkl. Validierung und Antwort). Im Admin-Einsätze-Komponent "
        "(src/components/Einsaetze.js) wurden die CRUD-Calls inklusive Fehlerbehandlung "
        "(try/catch + User-Feedback) ergänzt. Auf allen Speichern-Buttons sind jetzt sichtbare "
        "Erfolgs-/Fehlermeldungen sichtbar."
    ),
    where_to_test=[
        "Als Kunde im Dashboard: einen Termin öffnen, Uhrzeit ändern, 'Speichern' klicken.",
        "Als Admin im Einsätze-Tab: einen Einsatz öffnen, Datum verschieben, 'Speichern' klicken.",
        "Als Admin im Kundenprofil: ein Feld (z. B. Telefon) ändern, 'Speichern' klicken.",
        "Bei jedem dieser Schritte: die Seite neu laden und prüfen, ob die Änderung erhalten ist.",
    ],
    expected=(
        "Nach jedem Speichern erscheint ein Hinweis (grüne Bestätigung oder roter Fehlertext). "
        "Nach Neuladen sind die Änderungen tatsächlich in der DB persistiert."
    ),
    screenshot_caption=(
        "Vor- und Nach-Vergleich: Feld geändert, Speichern-Bestätigung sichtbar, Seite nach "
        "Reload mit dem neuen Wert."
    ),
)

add_item(
    doc, "F-24", "Kündigung innerhalb 30 Tagen löst CHF 300 Aufwandsentschädigung aus", "KRITISCH",
    problem=(
        "Kunden konnten den Vertrag kostenfrei innerhalb der ersten 30 Tage kündigen — entgegen "
        "den AGB und mit direktem finanziellen Schaden für PHC."
    ),
    change=(
        "src/pages/api/terminate-contract.js prüft jetzt die Tage seit Vertragsbeginn. Liegt "
        "die Kündigung unter 30 Tagen, wird via Stripe eine PaymentIntent über CHF 300 erstellt. "
        "Im Frontend (src/pages/dashboard/kundigung.js) erscheint ein Bestätigungsdialog mit "
        "expliziter Zustimmung zur Gebühr — ohne diese Zustimmung wird die Kündigung nicht "
        "verarbeitet."
    ),
    where_to_test=[
        "Erstellen Sie als Admin einen Test-Kunden mit einem Buchungsdatum vor weniger als 30 Tagen.",
        "Loggen Sie sich als dieser Kunde ein.",
        "Gehen Sie auf den Tab 'Kündigung' im Dashboard.",
        "Lesen Sie den Hinweis-Dialog inklusive Kostenangabe.",
        "Bestätigen Sie die Kündigung — beobachten Sie den Stripe-Charge im Backend (oder in "
        "Stripe-Dashboard → Payments).",
        "Wiederholen Sie den Test mit einem Kunden, dessen Buchung älter als 30 Tage ist — "
        "dort darf KEINE Gebühr erscheinen.",
    ],
    expected=(
        "Kunde innerhalb 30 Tagen: Dialog mit CHF 300-Hinweis, Stripe-Charge erfolgreich, "
        "Status wird auf 'gekuendigt' gesetzt. Kunde nach 30 Tagen: Dialog ohne Gebühr, keine "
        "Stripe-Buchung."
    ),
    screenshot_caption=(
        "Bestätigungsdialog mit CHF 300-Hinweis vor Kündigung + Stripe-Eintrag im Dashboard."
    ),
)

add_item(
    doc, "F-25", "Gekündigte Kunden bleiben mit Historie sichtbar", "KRITISCH",
    problem=(
        "Bei einer Kündigung wurde der Kundeneintrag stillschweigend aus dem Admin-Dashboard "
        "entfernt. Es gab keine Möglichkeit, gekündigte Kunden später nachzusehen, "
        "Kündigungsgrund oder -datum zu prüfen."
    ),
    change=(
        "Soft-Delete: bei Kündigung wird der User mit status='gekuendigt' inklusive "
        "terminationReason und terminationDate gespeichert (NICHT gelöscht). Im Admin gibt es "
        "eine neue Sektion /admin/ehemalige-kunden, die genau diese Liste anzeigt — mit Status, "
        "Grund und Datum. Aus der aktiven Kunden- und Buchungsübersicht werden sie korrekt "
        "ausgefiltert."
    ),
    where_to_test=[
        "Kündigen Sie als Kunde testweise einen Vertrag (siehe F-24).",
        "Loggen Sie sich als Admin ein und öffnen Sie im linken Menü 'Ehemalige Kunden'.",
        "Suchen Sie den eben gekündigten Kunden in der Liste.",
        "Prüfen Sie ob Status, Kündigungsgrund und Datum korrekt angezeigt werden.",
        "Öffnen Sie das Dashboard 'Übersicht' und stellen Sie sicher, dass dieser Kunde dort "
        "NICHT mehr in 'Neue Buchungen' erscheint.",
    ],
    expected=(
        "Ehemaliger Kunde ist in /admin/ehemalige-kunden vollständig dokumentiert. Aktive "
        "Dashboards zeigen ihn nicht mehr. Datenbank-Eintrag existiert weiter (Soft-Delete)."
    ),
    screenshot_caption=(
        "/admin/ehemalige-kunden-Tabelle mit Spalten Name, Status, Grund, Kündigungsdatum."
    ),
)

doc.add_page_break()

# ── Phase 2: E-Mail-Infrastruktur ────────────────────────────────────────

doc.add_heading("Phase 2 — E-Mail-Infrastruktur", level=1)
doc.add_paragraph(
    "Das E-Mail-System war im Audit als 'grösstes systemisches Problem' eingestuft: "
    "8+ hardcodierte Templates in mailer.js, kein Admin-Editor, Doppelversand, falsche "
    "Empfänger. Phase 2 hat die ganze Infrastruktur neu aufgesetzt."
)
doc.add_paragraph()

add_item(
    doc, "F-14 / F-43", "Alle E-Mail-Vorlagen sind in der DB und im Admin editierbar", "KRITISCH",
    problem=(
        "Templates waren im Code hardcodiert. Jede Textänderung brauchte ein Deployment. Im "
        "Admin-Dashboard fehlte der entsprechende Editor komplett."
    ),
    change=(
        "Neues Prisma-Modell EmailTemplate (name UNIQUE, subject, body, updatedAt). "
        "src/lib/emailTemplate.js → renderEmail() lädt zuerst aus der DB, fällt nur dann auf "
        "den hardcodierten Default zurück, wenn die Zeile fehlt. mailer.js wurde komplett "
        "refactored. Editor liegt unter /admin/email-templates. Seed: 11 Basis-Templates via "
        "`npm run seed:emails`."
    ),
    where_to_test=[
        "Als Admin /admin/email-templates öffnen.",
        "Eine Vorlage öffnen (z. B. 'appointmentConfirmation') und einen kleinen Test-Text "
        "in den Body einfügen — z. B. '[TEST 20.05]'.",
        "Speichern.",
        "Eine Testbuchung erstellen und die Bestätigungs-E-Mail auslösen.",
        "Im Postfach prüfen, ob '[TEST 20.05]' im Body sichtbar ist.",
        "Den Text in der Vorlage wieder entfernen.",
    ],
    expected=(
        "Änderungen im Admin-Editor werden ohne Deployment in der nächsten gesendeten E-Mail "
        "sichtbar. Beim Editor-Save bleibt der ursprüngliche Eintrag (Seed) erhalten und wird "
        "überschrieben."
    ),
    screenshot_caption=(
        "Editor /admin/email-templates mit geladener Vorlage + Beispiel-E-Mail mit dem "
        "geänderten Text im Postfach."
    ),
)

add_item(
    doc, "F-16", "Neuer Passwort-Flow für Erstanmeldung", "HOCH",
    problem=(
        "Bei Erstregistrierung bekam der/die Mitarbeiter:in eine 'Passwort zurücksetzen'-Mail. "
        "Das ist sprachlich falsch (man setzt zum ersten Mal eines, man setzt nichts zurück) "
        "und vermittelt den Eindruck eines bestehenden Accounts."
    ),
    change=(
        "Neue Route /set-password (eigenes Template, eigener Flow). Bei "
        "Mitarbeiter-Freischaltung wird ein 32-Byte resetToken (Gültigkeit 48 h) erstellt und "
        "in der Begrüssungs-E-Mail ('employeeWelcome') verlinkt. Klare Sprache: 'Willkommen, "
        "bitte erstellen Sie Ihr Passwort'."
    ),
    where_to_test=[
        "Als Admin einen neuen Mitarbeiter genehmigen (Status auf 'approved').",
        "Im Postfach des Mitarbeiters die Begrüssungs-E-Mail öffnen.",
        "Den 'Passwort erstellen'-Link klicken — er führt auf /set-password?email=...&token=...",
        "Ein Passwort setzen.",
        "Sich mit den neuen Daten einloggen.",
    ],
    expected=(
        "E-Mail-Wortlaut spricht von 'erstellen', nicht 'zurücksetzen'. Link ist 48 Stunden "
        "gültig. Login danach funktioniert."
    ),
    screenshot_caption=(
        "Begrüssungs-E-Mail im Postfach + die /set-password-Seite nach dem Klick."
    ),
)

add_item(
    doc, "F-15 / F-17", "Vorlagen-Texte korrekt + formelle Anrede", "HOCH",
    problem=(
        "Die ursprünglichen E-Mail-Texte stimmten nicht 1:1 mit den von PHC gelieferten "
        "Vorlagen überein. Die Anrede war zu direkt ('Hallo Maria') statt formell "
        "('Sehr geehrte Frau Müller')."
    ),
    change=(
        "src/lib/salutation.js → formalGreeting() rendert 'Sehr geehrter Herr {Name}' bzw. "
        "'Sehr geehrte Frau {Name}' aus den Anrede-Feldern. Alle 11 Templates wurden auf "
        "{{greeting}} umgestellt. Die endgültigen PHC-Marketing-Texte werden über /admin/"
        "email-templates eingepflegt — der Seed legt nur eine erste Arbeitsversion an."
    ),
    where_to_test=[
        "Sicherstellen, dass die endgültigen Marketing-Texte aus den Word-Dokumenten in /admin/"
        "email-templates eingefügt sind (das ist eine PHC-Aufgabe).",
        "Eine Testbuchung mit Anrede 'Frau' anlegen.",
        "Die Bestätigungs-E-Mail im Postfach öffnen.",
        "Prüfen, dass der Begrüssungssatz formell ist ('Sehr geehrte Frau ...').",
    ],
    expected=(
        "Begrüssung beginnt mit 'Sehr geehrte/r Herr/Frau {Nachname}'. Inhalt entspricht den "
        "von PHC gelieferten Marketing-Texten."
    ),
    screenshot_caption=(
        "E-Mail-Anfang im Postfach mit formeller Anrede + Editor-Ansicht der Vorlage."
    ),
)

doc.add_page_break()

# ── Phase 3: Datenintegrität & Anzeige ────────────────────────────────────

doc.add_heading("Phase 3 — Datenintegrität & Anzeige", level=1)
doc.add_paragraph(
    "Phase 3 hat die Datenstrukturen geradegezogen und die Anzeige im Dashboard so angepasst, "
    "dass alle gebuchten Informationen 1:1 sichtbar sind."
)
doc.add_paragraph()

add_item(
    doc, "F-01", "Dienstleistungen pro Wochentag wählbar", "KRITISCH",
    problem=(
        "Bei mehreren Schedule-Zeilen (z. B. Mo + Mi) musste der Kunde dieselben Subservices "
        "für beide Tage wählen oder gar keine. Eine differenzierte Auswahl ('Montag Haushalt, "
        "Mittwoch Spazieren') war nicht möglich."
    ),
    change=(
        "Die Subservice-Auswahl in src/pages/register-client.js wurde von einem globalen Array "
        "auf eine Per-Schedule-Struktur umgebaut. Jeder schedules[]-Eintrag trägt seine eigene "
        "Subservice-Liste, die State- und UI-Logik wurde entsprechend angepasst."
    ),
    where_to_test=[
        "Anonym /betreuung-zuhause-organisieren öffnen.",
        "Zwei Schedule-Zeilen anlegen (Mo + Mi).",
        "In Zeile 1 (Mo) Subservice 'Haushalt' wählen.",
        "In Zeile 2 (Mi) Subservice 'Spazieren' wählen.",
        "Buchung abschliessen und als der Kunde einloggen.",
        "Im Dashboard prüfen, ob die Tage tatsächlich die jeweiligen Subservices haben.",
    ],
    expected=(
        "Mo zeigt Haushalt, Mi zeigt Spazieren. Die Auswahl bleibt unabhängig pro Tag."
    ),
    screenshot_caption=(
        "Buchungsstrecke mit zwei Zeilen, jede mit eigener Subservice-Auswahl + danach das "
        "Kunden-Dashboard mit den getrennt angezeigten Diensten."
    ),
)

add_item(
    doc, "F-02", "Stundenberechnung bei mehreren Diensten korrekt", "KRITISCH",
    problem=(
        "Wenn der Kunde zwei Subservices wählte und einen wieder abwählte, blieb die "
        "Stundenanzahl falsch (z. B. 4 h / CHF 177 statt 2 h)."
    ),
    change=(
        "Die Berechnungslogik in src/pages/register-client.js läuft jetzt sauber über die "
        "aktuelle Subservice-Auswahl. Formel: Preis = Stunden × CHF 59. Aktualisierung bei "
        "jeder Änderung."
    ),
    where_to_test=[
        "Im Buchungsfunnel zwei Subservices wählen — der Preis aktualisiert sich.",
        "Einen der Subservices wieder abwählen.",
        "Stundenzahl und Preis-Anzeige prüfen.",
    ],
    expected=(
        "Stundenanzahl und Preis entsprechen exakt der aktuell ausgewählten Subservice-Liste."
    ),
    screenshot_caption=(
        "Buchungsstrecke mit Live-Berechnung: vor und nach Abwahl eines Subservice."
    ),
)

add_item(
    doc, "F-03 / F-28", "Hausnummer, Kanton, Anrede überall korrekt", "HOCH",
    problem=(
        "Die Hausnummer wurde nicht abgefragt und nirgends gespeichert. Kanton und Anrede "
        "tauchten teilweise leer in der Profilanzeige auf."
    ),
    change=(
        "Prisma-Schema wurde um houseNumber erweitert. Im Buchungsformular wird der Wert "
        "abgefragt. Im Kunden-Dashboard (personal-info.js, formular.js) werden Hausnummer, "
        "Kanton und Anrede aus der DB geladen und angezeigt — sowohl im View- als auch im "
        "Edit-Modus."
    ),
    where_to_test=[
        "Buchung mit Hausnummer '12a', Kanton 'ZH', Anrede 'Frau' abschliessen.",
        "Als dieser Kunde einloggen.",
        "Dashboard → Persönliche Info öffnen.",
        "Alle drei Felder sind sichtbar und korrekt.",
        "Auf 'Bearbeiten' klicken, ein Feld ändern, speichern, neu laden.",
    ],
    expected=(
        "Hausnummer, Kanton und Anrede sind in View und Edit identisch. Änderungen werden "
        "persistent gespeichert."
    ),
    screenshot_caption=(
        "Dashboard → Persönliche Info, View- und Edit-Modus nebeneinander."
    ),
)

add_item(
    doc, "F-34 / F-35", "Sprachen und Subservices vollständig übertragen", "HOCH",
    problem=(
        "Bei der Mitarbeiter-Registrierung gingen Felder wie languages, communicationTraits "
        "und dietaryExperience verloren. Subservices wurden zwischen Kunden- und Mitarbeiter-"
        "Seite unterschiedlich benannt."
    ),
    change=(
        "src/pages/api/employee-register.js speichert alle Formularfelder explizit. "
        "src/lib/serviceCatalog.js ist die zentrale Single-Source-of-Truth für Service- und "
        "Subservice-Namen — sowohl Kunden- als auch Mitarbeiter-Seite laden daraus."
    ),
    where_to_test=[
        "Bewerbung anlegen mit mehreren Sprachen, einer anderen Sprache (frei), "
        "Kommunikationsmerkmalen und Ernährungserfahrung.",
        "Als Admin den Kandidaten unter /admin/bewerber öffnen.",
        "Alle Felder müssen sichtbar sein.",
        "Im Subservice-Bereich prüfen, dass die Begriffe (z. B. 'kochen', 'spazieren') "
        "identisch sind mit denen auf der Kundenseite.",
    ],
    expected=(
        "Alle eingegebenen Bewerbungsfelder sind im Admin sichtbar. Subservice-Begriffe sind "
        "kongruent zwischen Kunde und Mitarbeiter."
    ),
    screenshot_caption=(
        "Admin-Detailansicht des Kandidaten mit allen ausgefüllten Feldern."
    ),
)

add_item(
    doc, "F-26", "Stornierung statt Rückerstattung — neue Logik", "HOCH",
    problem=(
        "Wording: 'Rückerstattung' war falsch — der Kunde soll eine Stornogebühr zahlen, nicht "
        "etwas zurückerstattet bekommen. Auch die Logik war umgekehrt: 7–14 Tage = 50 % "
        "Rückerstattung statt 50 % Stornogebühr."
    ),
    change=(
        "src/pages/client-dashboard.js → getCancellationFeeInfo() rechnet jetzt: ab 14 Tage = "
        "0 % Gebühr, 7–13 Tage = 50 % Gebühr, <7 Tage = 100 % Gebühr. Wording überall auf "
        "'Stornogebühr' umgestellt. Bestätigungs-E-Mail zeigt feePercent."
    ),
    where_to_test=[
        "Als Kunde einen Termin stornieren, der mehr als 14 Tage in der Zukunft liegt → "
        "0 % Gebühr.",
        "Termin in 10 Tagen stornieren → 50 % Gebühr.",
        "Termin in 3 Tagen stornieren → 100 % Gebühr.",
        "Wording im Dialog und in der Bestätigungs-E-Mail prüfen — überall steht 'Stornogebühr'.",
    ],
    expected=(
        "Gebühren-Schwellen stimmen. Wortlaut nutzt 'Stornogebühr', nicht 'Rückerstattung'."
    ),
    screenshot_caption=(
        "Stornierungs-Dialog mit angezeigter Gebühr + zugehörige Bestätigungs-E-Mail."
    ),
)

add_item(
    doc, "F-39", "Matching-Score: gewichtete Mitarbeiter-Empfehlung", "HOCH",
    problem=(
        "Der Admin musste passende Mitarbeiter manuell suchen — kein Score, keine Filter, "
        "keine Empfehlung."
    ),
    change=(
        "src/pages/api/admin/matchmaking.js implementiert jetzt eine gewichtete "
        "Score-Berechnung: Verfügbarkeit 30 %, Region/Distanz 25 %, Services 20 %, Sprache 10 %, "
        "Erfahrung 10 %, spezielle Anforderungen 5 %. Im Admin-UI erscheint bei 'Mitarbeiter "
        "zuweisen' eine Top-5-Liste mit Score-Balken und Aufschlüsselung."
    ),
    where_to_test=[
        "Als Admin einen offenen Einsatz öffnen.",
        "Auf 'Mitarbeiter zuweisen' klicken.",
        "Die Empfehlungsliste mit Score-Balken einsehen.",
        "Den Top-1-Kandidaten gegen die Buchung manuell prüfen: passt Region? Sprache? "
        "Verfügbarkeit?",
    ],
    expected=(
        "Top-5-Liste mit Prozent-Score, sortiert absteigend. Hover oder Klick zeigt die "
        "Aufschlüsselung pro Kriterium."
    ),
    screenshot_caption=(
        "Admin-Dialog 'Mitarbeiter zuweisen' mit sichtbaren Score-Balken."
    ),
)

add_item(
    doc, "F-10", "Betreuer erst sichtbar nach Annahme", "HOCH",
    problem=(
        "Der Kunde sah im Dashboard sofort den zugewiesenen Mitarbeiter — auch bevor dieser "
        "den Einsatz angenommen hatte. Bei Ablehnung war der Name dann plötzlich wieder weg."
    ),
    change=(
        "Im Kunden-Dashboard wird der Mitarbeiter nur angezeigt, wenn "
        "Assignment.confirmationStatus === 'confirmed' bzw. 'accepted'. Vorher steht "
        "'Mitarbeiter wird zugewiesen'."
    ),
    where_to_test=[
        "Als Admin einem Kunden einen Mitarbeiter zuweisen.",
        "Sofort als Kunde einloggen und Dashboard öffnen — Betreuer darf NICHT sichtbar sein.",
        "Als Mitarbeiter die Anfrage annehmen.",
        "Als Kunde Dashboard neu laden — Betreuer wird jetzt angezeigt.",
    ],
    expected=(
        "Mitarbeitername erscheint erst nach Annahme. Vorher Platzhalter-Text."
    ),
    screenshot_caption=(
        "Kunden-Dashboard vor Annahme (Platzhalter) + nach Annahme (Name sichtbar)."
    ),
)

add_item(
    doc, "F-11", "Workflow: Anfrage vs. Zuweisung getrennt", "MITTEL",
    problem=(
        "Im UI und in der Sprache war nicht klar getrennt: eine 'Zuweisung' wurde benutzt für "
        "etwas, das eigentlich erst eine 'Anfrage' an den Mitarbeiter war."
    ),
    change=(
        "Zwei-Schritte-Workflow: 1) Anfrage an MA, 2) erst nach Bestätigung = Zuweisung. "
        "Wording konsequent durchgezogen: im Admin steht 'Anfrage senden', im Kunden-Dashboard "
        "'Mitarbeiter wird zugewiesen', im Mitarbeiter-Dashboard 'Ausstehend' → 'Einsätze'."
    ),
    where_to_test=[
        "Den oben in F-10 beschriebenen Ablauf durchlaufen.",
        "Auf das genaue Wording in allen drei Rollen achten.",
    ],
    expected=(
        "'Anfrage' wird verwendet, solange der Mitarbeiter nicht bestätigt hat. Erst danach "
        "spricht das UI von 'Zuweisung' / 'Einsatz'."
    ),
    screenshot_caption=(
        "Drei nebeneinandergelegte Screenshots: Admin-Anfrage, Kunden-Pending, "
        "Mitarbeiter-Ausstehend."
    ),
)

doc.add_page_break()

# ── Phase 4: UI/UX Umbau ─────────────────────────────────────────────────

doc.add_heading("Phase 4 — UI/UX Umbau", level=1)
doc.add_paragraph(
    "Phase 4 betrifft die sichtbaren Strukturen: Menüs, Übersichten, Detailseiten. Hier wurden "
    "die Wünsche aus dem Audit (8-Punkte-Menü Kunde, 7-Punkte-Menü Mitarbeiter, "
    "Benachrichtigungs-Center etc.) konkret umgesetzt."
)
doc.add_paragraph()

add_item(
    doc, "F-27", "Kunden-Dashboard: neues Menü mit 8 Punkten", "HOCH",
    problem="Das alte Menü hatte eine andere Reihenfolge und einige Punkte fehlten.",
    change=(
        "navItems in src/pages/client-dashboard.js wurde umgebaut auf: Dashboard, Persönliche "
        "Info, Finanzen, Meine Verträge, Neue Buchung, Kündigung, Nachrichten & Feedback, "
        "Kontakt."
    ),
    where_to_test=[
        "Als Kunde einloggen.",
        "Die linke Navigation öffnen und durch alle 8 Punkte klicken.",
        "Reihenfolge prüfen (siehe oben).",
    ],
    expected=(
        "Die 8 Menüpunkte erscheinen genau in der genannten Reihenfolge. Jeder Klick führt "
        "auf eine funktionierende Seite."
    ),
    screenshot_caption=(
        "Linke Navigation im Kunden-Dashboard mit allen 8 Punkten sichtbar."
    ),
)

add_item(
    doc, "F-33", "Mitarbeiter-Dashboard: neues Menü", "HOCH",
    problem="Altes Menü entsprach nicht dem Konzept vom 26.01.",
    change=(
        "src/components/EmployeeLayout.js neu strukturiert: Dashboard, Einsätze "
        "(mit Sub-Tabs Kommend / Laufend / Vergangen), Zuweisungen, Abwesenheiten, "
        "Nachrichten, Profil, Dokumente, Logout."
    ),
    where_to_test=[
        "Als Mitarbeiter einloggen.",
        "Linke Navigation öffnen und alle Punkte abklicken.",
        "Im Punkt 'Einsätze' die drei Sub-Tabs anklicken: Kommend / Laufend / Vergangen.",
    ],
    expected=(
        "Alle Menüpunkte vorhanden und in der genannten Reihenfolge. Sub-Tabs filtern die "
        "Einsätze korrekt."
    ),
    screenshot_caption=(
        "Mitarbeiter-Dashboard mit linker Navigation + offene Einsätze-Sub-Tabs."
    ),
)

add_item(
    doc, "F-29", "Fragebogen 1:1 im Dashboard", "HOCH",
    problem=(
        "Der Fragebogen im Dashboard war unvollständig — Felder die in der Buchung erfasst "
        "wurden, fehlten in der Anzeige. Gleichzeitig waren überflüssige Felder vorhanden."
    ),
    change=(
        "src/pages/dashboard/formular.js wurde komplett überarbeitet: alle bei der Buchung "
        "abgefragten Felder werden 1:1 angezeigt. View- und Edit-Modus haben dieselbe "
        "Feld-Liste. Überflüssige Felder entfernt."
    ),
    where_to_test=[
        "Eine Buchung mit möglichst vielen ausgefüllten Feldern abschliessen "
        "(Gesundheit, Mobilität, Hilfsmittel, Sprachen etc.).",
        "Als dieser Kunde im Dashboard → Formular öffnen.",
        "Mit der ursprünglichen Buchung vergleichen — jedes Feld muss vorhanden sein.",
    ],
    expected=(
        "Jedes bei der Buchung erfasste Feld erscheint im Fragebogen-View."
    ),
    screenshot_caption=(
        "Dashboard → Formular, vollständige Feld-Liste im View-Modus."
    ),
)

add_item(
    doc, "F-40", "Benachrichtigungs-Center im Admin", "HOCH",
    problem="Admin hatte keinen Überblick über offene Anfragen, neue Kandidaten etc.",
    change=(
        "src/components/AdminLayout.js zeigt jetzt oben rechts ein Glocken-Icon mit Badge "
        "(ungelesene Anzahl). Klick öffnet das Notification-Panel. API-Quelle: "
        "src/pages/api/admin/notifications.js."
    ),
    where_to_test=[
        "Als Admin im Dashboard die Glocke oben rechts beachten.",
        "Eine neue Bewerbung von einem Test-Account aus einreichen.",
        "Im Admin: Badge-Zähler erhöht sich.",
        "Glocke anklicken — Liste enthält den neuen Kandidaten.",
    ],
    expected="Badge zeigt korrekte ungelesene Anzahl. Panel listet die Items.",
    screenshot_caption=(
        "Glocken-Icon mit Badge + geöffnetes Notification-Panel."
    ),
)

add_item(
    doc, "F-42", "Einsätze: CRUD, Filter, Verlinkungen funktionieren", "KRITISCH",
    problem=(
        "Im Admin-Tab 'Einsätze' funktionierte Speichern nicht. Es fehlten Filter (Status, "
        "Datum, Mitarbeiter) und Verlinkungen zu Kunden- und Mitarbeiter-Profilen."
    ),
    change=(
        "src/components/Einsaetze.js überarbeitet: Filter nach Status/Datum/MA/Kunde, "
        "klickbare Links zu Profilseiten, CRUD über /api/admin/schedules/[id]/edit. "
        "Status-Wechsel inkl. 'storniert'."
    ),
    where_to_test=[
        "Als Admin /admin/einsaetze öffnen.",
        "Filter testen: nach Status (z. B. 'angefragt'), nach Datum, nach Mitarbeiter.",
        "Einen Einsatz öffnen, Datum ändern, speichern, neu laden — Änderung bleibt.",
        "Auf den Kundennamen klicken → Kundenprofil öffnet sich.",
        "Auf den Mitarbeiternamen klicken → Mitarbeiterprofil öffnet sich.",
    ],
    expected=(
        "Filter funktionieren. Speichern persistiert. Beide Verlinkungen funktionieren."
    ),
    screenshot_caption=(
        "/admin/einsaetze mit aktivem Filter und sichtbaren Verlinkungen."
    ),
)

add_item(
    doc, "F-22", "Subservices vollständig im Kunden-Dashboard", "HOCH",
    problem="Manche gebuchten Subservices wurden im Dashboard nicht angezeigt.",
    change=(
        "Die API liefert jetzt alle Subservices aus Schedule und User. AssignmentList.js und "
        "client-dashboard.js zeigen sie vollständig in der Vertrags- und Einsatz-Ansicht."
    ),
    where_to_test=[
        "Buchung mit mehreren Subservices abschliessen.",
        "Im Dashboard → Meine Verträge prüfen ob alle Subservices erscheinen.",
    ],
    expected="Alle bei der Buchung gewählten Subservices sind sichtbar.",
    screenshot_caption=(
        "Dashboard → Meine Verträge mit vollständiger Subservice-Liste."
    ),
)

add_item(
    doc, "F-32", "Nächster Einsatz mit allen Details (Adresse, Subservices, Dauer)", "HOCH",
    problem="Im Mitarbeiter-Dashboard fehlten beim nächsten Einsatz Adresse, Subservices und Dauer.",
    change=(
        "src/components/AssignmentList.js zeigt jetzt clientStreet, clientPLZ, clientCity, "
        "subServices, baseHours und baseKm."
    ),
    where_to_test=[
        "Als Mitarbeiter einloggen.",
        "Im Dashboard die 'Nächster Einsatz'-Karte und die Einsätze-Liste anschauen.",
        "Adresse, Subservices, Stundenzahl, Einsatzort prüfen.",
    ],
    expected="Alle Detailfelder sind sichtbar. Adresse wird vollständig angezeigt.",
    screenshot_caption=(
        "Mitarbeiter-Dashboard, Karte 'Nächster Einsatz' mit allen Details."
    ),
)

add_item(
    doc, "F-44", "Bewerber → Kandidaten + erweitertes Profil", "MITTEL",
    problem=(
        "Das Wort 'Bewerber' sollte im UI durch 'Kandidaten' ersetzt werden. Im "
        "Detail-Profil fehlten Geburtsdatum, Raucher, Kanton, PLZ."
    ),
    change=(
        "AdminLayout-Label heisst jetzt 'Kandidaten' (Route /admin/bewerber bleibt aus "
        "Kompatibilitätsgründen). Komponente intern in KandidatenPage umbenannt. Die "
        "Detailseite /admin/employees/[id] zeigt zusätzlich Geburtsdatum, Raucher-Status, "
        "Kanton, PLZ."
    ),
    where_to_test=[
        "Als Admin im Menü links sollte der Eintrag 'Kandidaten' (NICHT 'Bewerber') heissen.",
        "Klick öffnet die Liste.",
        "Einen Kandidaten anklicken — im Detail Geburtsdatum, Raucher, Kanton, PLZ prüfen.",
    ],
    expected="Label heisst 'Kandidaten'. Detail-Profil enthält alle 4 zusätzlichen Felder.",
    screenshot_caption=(
        "Admin-Menü mit 'Kandidaten'-Eintrag + Detail-Profil mit allen Feldern."
    ),
)

add_item(
    doc, "F-45", "Mitarbeiter-Profil: Verfügbarkeit, DL, Reisebereitschaft korrekt", "HOCH",
    problem="Verfügbarkeitsanzeige war unübersichtlich, Reisebereitschaft inkonsistent mit Fragebogen.",
    change=(
        "/admin/employees/[id] zeigt availabilityFrom + availabilityDays strukturiert. "
        "Reisebereitschaft (travelSupport) wird kongruent zum Bewerbungsformular dargestellt. "
        "Dienstleistungen übersichtlicher gruppiert."
    ),
    where_to_test=[
        "Als Admin einen Mitarbeiter öffnen.",
        "Verfügbarkeit prüfen (Tage + Zeiten).",
        "Reisebereitschaft mit der Bewerbung vergleichen.",
        "Dienstleistungs-Liste anschauen.",
    ],
    expected="Alle drei Bereiche sind klar und stimmen mit den Originalfeldern aus der Bewerbung überein.",
    screenshot_caption=(
        "Admin → Mitarbeiter-Detail, alle drei Bereiche sichtbar."
    ),
)

add_item(
    doc, "F-36", "Mitarbeiter-Dashboard aktualisiert sich automatisch", "HOCH",
    problem=(
        "Eine vom Admin geschickte Anfrage erschien beim Mitarbeiter erst nach Re-Login. "
        "Manchmal wurde dieselbe Zuweisung zweifach angezeigt."
    ),
    change=(
        "src/pages/employee-dashboard.js refetcht jetzt automatisch: bei jedem Tab-Focus, "
        "bei Visibility-Change und alle 60 Sekunden im Hintergrund (gated auf "
        "visibilityState='visible'). reloadAssignments() dedupliziert per ID."
    ),
    where_to_test=[
        "Als Mitarbeiter einloggen und Dashboard offen lassen.",
        "Parallel als Admin eine neue Anfrage an genau diesen Mitarbeiter senden.",
        "Im Mitarbeiter-Tab: ohne Re-Login warten — die neue Anfrage erscheint spätestens "
        "nach 60 Sekunden, schneller wenn man den Tab kurz wechselt und zurückkehrt.",
        "Keine Duplikate in der Liste.",
    ],
    expected=(
        "Anfrage taucht spätestens nach 60 s automatisch auf. Liste enthält jeden Eintrag genau einmal."
    ),
    screenshot_caption=(
        "Vorher/Nachher: Dashboard leer → Anfrage erscheint."
    ),
)

add_item(
    doc, "F-20", "Status 'zugewiesen' korrekt nach Annahme", "HOCH",
    problem="Status blieb auf 'wird organisiert' obwohl ein Mitarbeiter bereits zugewiesen war.",
    change="Status-Mapping in client-dashboard.js berücksichtigt jetzt confirmationStatus.",
    where_to_test=[
        "Den Annahme-Flow aus F-10 ausführen.",
        "Im Kunden-Dashboard das Status-Label prüfen.",
    ],
    expected="Nach Annahme zeigt das Label 'Zugewiesen' bzw. 'Bestätigt'.",
    screenshot_caption="Kunden-Dashboard mit korrektem Status-Label.",
)

add_item(
    doc, "F-31", "Mitarbeiter-Session-Timeout: sauberer Logout", "HOCH",
    problem="Bei abgelaufenem Token zeigte das Dashboard leere Listen statt sich auszuloggen.",
    change=(
        "authFetch() in src/lib/clientAuth.js erkennt 401-Antworten und leitet sofort auf "
        "/login um. localStorage wird geleert."
    ),
    where_to_test=[
        "Als Mitarbeiter einloggen.",
        "In den DevTools den 'userToken' aus localStorage löschen.",
        "Eine beliebige Aktion ausführen (z. B. Tab wechseln).",
    ],
    expected="Sofortige Umleitung auf /login.",
    screenshot_caption="DevTools mit gelöschtem Token + Browser auf /login.",
)

doc.add_page_break()

# ── Phase 5: Feinschliff ─────────────────────────────────────────────────

doc.add_heading("Phase 5 — Feinschliff & Wording", level=1)
doc.add_paragraph(
    "Phase 5 deckt Sprache, Cookies, Reminder-Cron und kleinere Polish-Punkte ab. Diese "
    "Punkte sind nicht kritisch, aber wichtig für die Wahrnehmung als professionelles System."
)
doc.add_paragraph()

add_item(
    doc, "F-07", "Wording 'Buchung' statt 'Registrierung' (kundenseitig)", "MITTEL",
    problem=(
        "Kunden sahen 'Registrierung erfolgreich!' — sprachlich falsch, weil ein Kunde bucht, "
        "kein Konto registriert."
    ),
    change=(
        "src/pages/register.js zeigt jetzt rollen-abhängig: Kunden sehen 'Buchung "
        "erfolgreich!', Bewerber sehen 'Bewerbung erfolgreich!'."
    ),
    where_to_test=[
        "Anonym /register öffnen.",
        "Rolle 'Client' wählen, Formular ausfüllen, einreichen.",
        "Erfolgsmeldung lesen.",
        "Den Test mit Rolle 'Employee' wiederholen.",
    ],
    expected="Client → 'Buchung erfolgreich!'. Employee → 'Bewerbung erfolgreich!'.",
    screenshot_caption="Erfolgsseite jeweils für Client- und Employee-Rolle.",
)

add_item(
    doc, "F-37", "Englische Begriffe übersetzt", "HOCH",
    problem="In Dashboards tauchten Status wie 'pending', 'approved', 'rejected' auf.",
    change=(
        "src/lib/statusLabels.js zentralisiert alle Übersetzungen. labelFor() wird in allen "
        "JSX-Stellen verwendet. Raw-Status erscheinen nirgends mehr."
    ),
    where_to_test=[
        "Alle drei Dashboards (Admin, Kunde, Mitarbeiter) durchklicken.",
        "Auf englische Begriffe achten — es darf keine geben.",
    ],
    expected="Alle Status-Labels sind deutsch.",
    screenshot_caption="Dashboard mit sichtbaren Status-Chips/Badges in deutscher Sprache.",
)

add_item(
    doc, "F-47", "Cookie-Banner + Datenschutz", "MITTEL",
    problem="Es gab keinen Cookie-Hinweis. DSG/DSGVO-Verstoss.",
    change=(
        "src/components/CookieBanner.js wird in _app.js eingebunden. Speichert Einwilligung "
        "in localStorage (Schlüssel 'phc_cookie_consent_v1'). Verlinkt /datenschutz. "
        "Standard: nur notwendige Cookies."
    ),
    where_to_test=[
        "Browser im Inkognito-Modus öffnen, auf die Startseite gehen.",
        "Cookie-Banner muss erscheinen.",
        "Auf 'Nur notwendige' klicken — Banner verschwindet, localStorage enthält Einwilligung.",
        "In einem neuen Inkognito-Fenster wiederholen mit 'Alle akzeptieren'.",
    ],
    expected="Banner erscheint beim ersten Besuch, verschwindet nach Auswahl, kommt nicht wieder.",
    screenshot_caption="Cookie-Banner sichtbar am Seitenrand.",
)

add_item(
    doc, "F-46", "Cron-Jobs: 2-Tage-Reminder + 30-Tage-Ablehnung", "MITTEL",
    problem=(
        "Nach Einladung zum Interview gab es keine automatische Erinnerung. Nach 30 Tagen "
        "ohne Reaktion blieb der Status hängen."
    ),
    change=(
        "src/pages/api/cron/interview-reminders.js: täglicher Job (vercel.json: '0 8 * * *'). "
        "Nach 2 Tagen ohne Interview-Buchung → Reminder. Nach 30 Tagen → automatischer Status "
        "auf 'rejected'."
    ),
    where_to_test=[
        "Test ist datumsabhängig — am besten in einer Stage-Umgebung das eingeladen-Datum "
        "manuell auf vor 3 Tagen setzen.",
        "Cron manuell triggern: GET /api/cron/interview-reminders.",
        "Postfach des Bewerbers auf Reminder-E-Mail prüfen.",
        "Test mit Datum vor 31 Tagen wiederholen — Status muss auf 'rejected' springen.",
    ],
    expected="Reminder kommt an Tag 2. Auto-Reject an Tag 30.",
    screenshot_caption="Cron-Log + zugehörige E-Mail im Postfach.",
)

add_item(
    doc, "F-04", "Hilfsmittel werden nur einmal angezeigt", "MITTEL",
    problem=(
        "Im Fragebogen wurden Hilfsmittel zweifach angezeigt, weil mehrere überlappende Felder "
        "(aids, mobilityAids, toolsAvailable etc.) separat gerendert wurden."
    ),
    change=(
        "src/pages/dashboard/formular.js nutzt dedupeCsv() — 5 Felder werden zu 2 zusammengeführt: "
        "'Hilfsmittel' und 'Sonstige Hilfsmittel'."
    ),
    where_to_test=[
        "Buchung mit mehreren Hilfsmitteln (z. B. Rollator, Gehstock) abschliessen.",
        "Im Dashboard → Formular die Sektion 'Gesundheit & Mobilität' öffnen.",
        "Hilfsmittel zählen — keine Duplikate erlaubt.",
    ],
    expected="Jedes Hilfsmittel erscheint genau einmal.",
    screenshot_caption="Dashboard → Formular, Sektion Hilfsmittel.",
)

add_item(
    doc, "F-21", "4 Termine + 'Alle anzeigen'", "MITTEL",
    problem="Es wurden nur 3 Termine angezeigt, ohne Möglichkeit alle zu sehen.",
    change=(
        "client-dashboard.js: TERMINE_PAGE_SIZE = 4. Button 'Alle anzeigen ({count})' "
        "klappt die volle Liste auf."
    ),
    where_to_test=[
        "Als Kunde mit mehr als 4 zukünftigen Terminen einloggen.",
        "Dashboard öffnen — 4 Termine sichtbar + Button.",
        "Button klicken — alle Termine sichtbar.",
    ],
    expected="4 standardmässig, Klick zeigt alle.",
    screenshot_caption="Termine-Liste vorher (4) und nachher (alle).",
)

add_item(
    doc, "F-30", "Zahlungsmethode lesbar + Gutschein-Sektion", "MITTEL",
    problem="Zahlungsmethode-UI war schlecht lesbar. Gutscheine fehlten komplett.",
    change=(
        "src/pages/dashboard/finanzen.js: Karten-Editor + Gutschein-Eingabe (handleVoucherSubmit "
        "→ /api/vouchers/use)."
    ),
    where_to_test=[
        "Als Kunde im Dashboard → Finanzen öffnen.",
        "Zahlungsmethode prüfen (lesbare Kartenmaske).",
        "Gutscheincode eingeben (z. B. Test-Code 'PHC50').",
    ],
    expected="Karten-Daten klar lesbar. Gutschein wird angenommen oder mit Fehlermeldung abgelehnt.",
    screenshot_caption="Dashboard → Finanzen mit Karten-UI + Gutschein-Eingabe.",
)

add_item(
    doc, "F-05", "Arbeitspensum: Mehrfachauswahl bei Bewerbung", "HOCH",
    problem="Bewerber konnten nur ein Arbeitspensum auswählen, obwohl mehrere realistisch sind.",
    change=(
        "Employee-Schema: desiredWeeklyHours von String auf String[] geändert. Formular nutzt "
        "Mehrfach-Checkbox."
    ),
    where_to_test=[
        "Bewerbung anlegen.",
        "Im Schritt 'Arbeitspensum' mehrere Optionen ankreuzen.",
        "Bewerbung einreichen.",
        "Als Admin den Kandidaten öffnen und alle gewählten Pensa prüfen.",
    ],
    expected="Mehrfachauswahl möglich. Alle Werte gespeichert.",
    screenshot_caption="Bewerbungsformular mit Mehrfach-Checkboxen.",
)

add_item(
    doc, "F-08", "Separate Adresse der anfragenden Person", "HOCH",
    problem=(
        "Die Adresse der anfragenden Person (Auftraggeber) wurde nicht separat erfasst — "
        "Hausnummer, Strasse, PLZ, Ort, Kanton fehlten."
    ),
    change=(
        "Prisma-Schema erweitert um requestStreet, requestHouseNumber, requestPostalCode, "
        "requestCity, requestKanton. Felder werden in der Buchungsstrecke erfragt."
    ),
    where_to_test=[
        "Buchung anlegen, bei der die anfragende Person eine andere Adresse hat als die "
        "betreute Person.",
        "Als Admin den Kunden öffnen und beide Adressen prüfen.",
    ],
    expected="Beide Adressen sind separat erfasst und sichtbar.",
    screenshot_caption="Admin-Kundendetail mit beiden Adressblöcken.",
)

add_item(
    doc, "F-19", "3× selber Tag — beseitigt durch F-09", "KRITISCH",
    problem=(
        "Im Dashboard standen 3× dasselbe Datum, obwohl der Kunde verschiedene Wochentage "
        "gebucht hatte. Folgefehler aus F-09."
    ),
    change=(
        "Mit dem F-09-Fix in register-client.js werden alle Schedule-Zeilen mit ihrem "
        "eigenen Wochentag persistiert. F-19 verschwindet automatisch."
    ),
    where_to_test=[
        "Den Test aus F-09 durchführen.",
        "Im Kunden-Dashboard die ersten 6 Termine anschauen.",
        "Drei verschiedene Wochentage sollten abwechselnd erscheinen.",
    ],
    expected="Termine fallen auf unterschiedliche Wochentage — keine Wiederholung desselben Datums.",
    screenshot_caption="Termin-Liste mit 6 sichtbaren Einträgen, unterschiedliche Daten.",
)

add_item(
    doc, "F-41", "Gekündigte Kunden nicht mehr bei neuen Buchungen", "HOCH",
    problem="Gekündigte Kunden tauchten weiter in der 'Neue Buchungen'-Liste auf.",
    change=(
        "src/pages/api/admin/dashboard-overview.js: WHERE-Klausel filtert "
        "status NOT IN ('gekuendigt', 'canceled', 'cancelled')."
    ),
    where_to_test=[
        "Einen Kunden kündigen (siehe F-24).",
        "Als Admin im Dashboard 'Übersicht' die Liste 'Neue Buchungen' prüfen — der gekündigte "
        "Kunde darf nicht erscheinen.",
        "/admin/ehemalige-kunden öffnen — dort ist er sichtbar.",
    ],
    expected="Gekündigter Kunde nur in /admin/ehemalige-kunden, nicht in der Übersicht.",
    screenshot_caption="Admin-Dashboard 'Neue Buchungen' ohne den gekündigten Kunden.",
)

doc.add_page_break()

# ── Offene Punkte ────────────────────────────────────────────────────────

doc.add_heading("Offene Punkte (keine Code-Aufgaben)", level=1)
doc.add_paragraph(
    "Drei Punkte sind aus Code-Sicht erledigt, brauchen aber noch eine Aktion von PHC bzw. "
    "vom QA-Team."
)

doc.add_heading("F-15 — Finale E-Mail-Texte einpflegen", level=2)
doc.add_paragraph(
    "Die 11 Templates sind in der DB als erste Arbeitsversion vorhanden. PHC (Marketing / "
    "Legal) muss die endgültigen Texte aus den Word-Dokumenten in /admin/email-templates "
    "einpflegen. Der Seed läuft idempotent — er überschreibt KEINE vom Admin gespeicherte "
    "Version."
)
doc.add_paragraph(
    "Verantwortlich: PHC. Aufwand: ca. 1–2 Stunden Copy-Paste."
)

doc.add_heading("F-19 — Manuelle Smoke-Test Buchung", level=2)
doc.add_paragraph(
    "Wir haben den F-09-Fix per Code-Review verifiziert. Eine manuelle Test-Buchung "
    "(Mo/Mi/Fr mit Beginn an einem Dienstag) ist die finale Bestätigung. Empfehlung: vor "
    "dem Soft-Launch durchspielen."
)

doc.add_heading("F-04 — Visuelle Stichprobe", level=2)
doc.add_paragraph(
    "Die Deduplizierung der Hilfsmittel-Felder ist im Code aktiv (dedupeCsv). Eine kurze "
    "Sicht-Prüfung im Dashboard ist trotzdem sinnvoll."
)

doc.add_page_break()


# ── Sign-off Checkliste ─────────────────────────────────────────────────

doc.add_heading("Sign-off-Checkliste", level=1)
doc.add_paragraph(
    "Diese Checkliste dient als Abnahmeprotokoll. Jeder Punkt wird durch Bettina oder Silvain "
    "abgehakt. Wenn alle Punkte erledigt sind und keine kritischen Bugs auftauchen, kann der "
    "Soft-Launch mit 3–5 Testkunden am 10.06.2026 starten."
)
doc.add_paragraph()

checklist_items = [
    "Phase 1: F-12 E-Mail-Empfänger getestet",
    "Phase 1: F-38 Session-Verhalten getestet",
    "Phase 1: F-09 Wochentage in der Buchung korrekt",
    "Phase 1: F-13/18 Mitarbeiter-E-Mail nur einmal",
    "Phase 1: F-06 Bewerbung endet auf Bestätigungsseite",
    "Phase 1: F-48/23/42 Speichern funktioniert überall",
    "Phase 1: F-24 CHF 300-Dialog bei Frühkündigung",
    "Phase 1: F-25 Ehemalige Kunden sichtbar",
    "Phase 2: F-14/43 Editor /admin/email-templates funktioniert",
    "Phase 2: F-16 Passwort-erstellen-Flow getestet",
    "Phase 2: F-15/17 finale Texte + formelle Anrede eingepflegt",
    "Phase 3: F-01 Subservices pro Tag wählbar",
    "Phase 3: F-02 Stundenberechnung korrekt",
    "Phase 3: F-03/28 Hausnummer/Kanton/Anrede überall",
    "Phase 3: F-34/35 Sprachen + Subservices übertragen",
    "Phase 3: F-26 Stornogebühr-Logik korrekt",
    "Phase 3: F-39 Matching-Score sichtbar",
    "Phase 3: F-10 Betreuer erst nach Annahme sichtbar",
    "Phase 3: F-11 Anfrage-/Zuweisungs-Wording",
    "Phase 4: F-27 Kunden-Menü 8 Punkte",
    "Phase 4: F-33 Mitarbeiter-Menü",
    "Phase 4: F-29 Fragebogen vollständig",
    "Phase 4: F-40 Benachrichtigungs-Center",
    "Phase 4: F-42 Einsätze CRUD",
    "Phase 4: F-22 Subservices vollständig",
    "Phase 4: F-32 Einsatz-Details",
    "Phase 4: F-44 Kandidaten-Umbenennung",
    "Phase 4: F-45 Mitarbeiter-Profil korrekt",
    "Phase 4: F-36 Dashboard-Auto-Refresh",
    "Phase 4: F-20 Status nach Annahme",
    "Phase 4: F-31 Session-Logout Mitarbeiter",
    "Phase 5: F-07 Buchungs-Wording",
    "Phase 5: F-37 Deutsche Status-Labels",
    "Phase 5: F-47 Cookie-Banner",
    "Phase 5: F-46 Cron-Reminder + Ablehnung",
    "Phase 5: F-04 Hilfsmittel dedupliziert",
    "Phase 5: F-21 4 Termine + Button",
    "Phase 5: F-30 Finanzen + Gutschein",
    "Phase 5: F-05 Mehrfach-Arbeitspensum",
    "Phase 5: F-08 Separate Auftraggeber-Adresse",
    "Phase 5: F-19 Smoke-Test bestanden",
    "Phase 5: F-41 Gekündigte ausgefiltert",
]

ctbl = doc.add_table(rows=1, cols=3)
ctbl.style = "Light Grid Accent 1"
hdr = ctbl.rows[0].cells
hdr[0].text = "Punkt"
hdr[1].text = "Tester"
hdr[2].text = "Datum / Bemerkung"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
for item in checklist_items:
    row = ctbl.add_row().cells
    row[0].text = "☐  " + item
    row[1].text = ""
    row[2].text = ""

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Freigabe Soft-Launch (10.06.2026): ").bold = True
p.add_run("☐ ja  ☐ nein").italic = True
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Unterschrift Bettina: _______________________     Datum: ____________")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Unterschrift Silvain: _______________________     Datum: ____________")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Unterschrift Fisnik:  _______________________     Datum: ____________")


# ── Save ─────────────────────────────────────────────────────────────────

import sys
out_path = "PHC_Test_Dokumentation.docx"
try:
    doc.save(out_path)
    print(f"Saved: {out_path}")
except PermissionError:
    # File is open in Word — write a versioned copy instead
    out_path = "PHC_Test_Dokumentation_v2.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}  (original was locked — close it in Word and re-run to overwrite)")
