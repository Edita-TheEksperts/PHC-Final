"""
Generates PHC_Platform_Documentation.docx — a single English Word document
covering every email the system sends and every fully-working feature of
the Prime Home Care platform, grouped by user role.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x0B, 0x4F, 0x6C)
ACCENT = RGBColor(0x1A, 0x76, 0xD2)
GREY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
ORANGE = RGBColor(0xE6, 0x5A, 0x00)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=BRAND):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 0:
        run.font.size = Pt(26)
    elif level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_kv_table(doc, rows, col_widths=(Cm(5), Cm(11.5))):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.text = ""
        c2.text = ""
        r1 = c1.paragraphs[0].add_run(k)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(v)
        r2.font.size = Pt(10)
        c1.width = col_widths[0]
        c2.width = col_widths[1]
        set_cell_bg(c1, "EAF2F8")
    return table


def add_emails_table(doc, emails):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    headers = ["#", "Email", "Trigger / When it's sent", "Recipient"]
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "0B4F6C")
    for idx, em in enumerate(emails, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = em["name"]
        row[2].text = em["trigger"]
        row[3].text = em["recipient"]
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


def add_features_table(doc, features):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Feature", "What the user can do", "Status"]):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "0B4F6C")
    for f in features:
        row = table.add_row().cells
        row[0].text = f["name"]
        row[1].text = f["desc"]
        row[2].text = f["status"]
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


def add_screenshot_placeholder(doc, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ Screenshot placeholder — {label} ]")
    run.italic = True
    run.font.color.rgb = GREY
    run.font.size = Pt(10)


# ----------------------------------------------------------------------
# DATA
# ----------------------------------------------------------------------

EMAILS = [
    {
        "name": "Client Welcome — \"Willkommen bei Prime Home Care\"",
        "trigger": "A new client completes registration and verifies email.",
        "recipient": "Client",
    },
    {
        "name": "Employee Welcome — \"Willkommen im Prime Home Care Team\"",
        "trigger": "A new employee completes registration; portal access is activated.",
        "recipient": "Employee",
    },
    {
        "name": "Application Received — \"Ihre Bewerbung bei Prime Home Care AG\"",
        "trigger": "Caregiver submits the employee application form.",
        "recipient": "Applicant (Employee)",
    },
    {
        "name": "Interview Invitation — \"Terminvereinbarung für ein persönliches Kennenlernen\"",
        "trigger": "Application received; system sends a Calendly link to book an interview.",
        "recipient": "Applicant (Employee)",
    },
    {
        "name": "Interview Reminder",
        "trigger": "Cron-based reminder before a scheduled interview if no booking has been made.",
        "recipient": "Applicant (Employee)",
    },
    {
        "name": "Approval Email + Rahmenvereinbarung (PDF attached)",
        "trigger": "Admin approves an employee application.",
        "recipient": "Employee",
    },
    {
        "name": "New Assignment Proposal — \"Neuer Einsatz – Bitte bestätigen\"",
        "trigger": "Admin assigns a job; employee has 36h to accept or reject.",
        "recipient": "Employee",
    },
    {
        "name": "12-hour Assignment Reminder",
        "trigger": "Cron, 12 hours after the proposal if no response yet.",
        "recipient": "Employee",
    },
    {
        "name": "24-hour Assignment Reminder (final)",
        "trigger": "Cron, 24 hours after the proposal — last call before auto-cancel.",
        "recipient": "Employee",
    },
    {
        "name": "Assignment Auto-Cancelled (admin notice)",
        "trigger": "Cron auto-cancels an unanswered proposal after 36h.",
        "recipient": "Admin",
    },
    {
        "name": "Caregiver Assigned — \"Ihr Betreuer wurde zugewiesen\"",
        "trigger": "Employee accepts the assignment.",
        "recipient": "Client",
    },
    {
        "name": "Assignment Contract (Einsatz-Arbeitsvertrag PDF)",
        "trigger": "Assignment confirmed — individual contract generated and emailed.",
        "recipient": "Employee",
    },
    {
        "name": "Capacity Limit — \"Information zur Verfügbarkeit Ihrer Buchung\"",
        "trigger": "Daily cron — no caregiver capacity for a booked period.",
        "recipient": "Client",
    },
    {
        "name": "Unassigned-Client Notice",
        "trigger": "Cron detects a client with a service but no caregiver yet.",
        "recipient": "Client",
    },
    {
        "name": "4-hour Service Reminder — \"Ihre Betreuung wartet auf Sie\"",
        "trigger": "4 hours before the scheduled appointment.",
        "recipient": "Client",
    },
    {
        "name": "48-hour Profile Reminder",
        "trigger": "48 hours before appointment — prompts client to complete missing info.",
        "recipient": "Client",
    },
    {
        "name": "Cancellation Confirmation",
        "trigger": "Client cancels a booking — email shows refund tier (100% / 80% / 50%).",
        "recipient": "Client",
    },
    {
        "name": "Assignment Cancelled (party-specific)",
        "trigger": "Any party cancels — separate emails to client, employee, admin.",
        "recipient": "Client + Employee + Admin",
    },
    {
        "name": "Payment Confirmation",
        "trigger": "Stripe charge succeeds.",
        "recipient": "Client",
    },
    {
        "name": "Payment Update Request",
        "trigger": "Client requests to change payment method/amount.",
        "recipient": "Admin",
    },
    {
        "name": "Feedback Request — \"Wie war Ihre Betreuung?\"",
        "trigger": "After service completion (manual or cron).",
        "recipient": "Client",
    },
    {
        "name": "Rejection Warning",
        "trigger": "Employee rejects 2 or more assignments — written warning.",
        "recipient": "Employee",
    },
    {
        "name": "Client Termination",
        "trigger": "Admin processes contract end (Abschluss form).",
        "recipient": "Client",
    },
    {
        "name": "Password Reset — \"Passwort zurücksetzen\"",
        "trigger": "User clicks Forgot Password — link valid 1 hour.",
        "recipient": "Client / Employee",
    },
    {
        "name": "System Maintenance Notice — \"Wartungsmitteilung\"",
        "trigger": "Admin triggers maintenance announcement to all clients.",
        "recipient": "All Clients",
    },
    {
        "name": "Contact-Form / Support Inquiry",
        "trigger": "Public contact form submitted on the website.",
        "recipient": "Admin / Support",
    },
    {
        "name": "Document Send-out",
        "trigger": "Admin manually sends a document (contract, form, PDF) to a user.",
        "recipient": "Client or Employee",
    },
    {
        "name": "Newsletter Welcome (Mailchimp)",
        "trigger": "Client opts into newsletter at registration.",
        "recipient": "Subscriber",
    },
]

ADMIN_FEATURES = [
    {"name": "Dashboard Overview",
     "desc": "Live KPIs: pending approvals, active clients, employees, open assignments, vacations, due tasks.",
     "status": "Live"},
    {"name": "Pending Approvals",
     "desc": "Review new client and employee registrations; approve or reject with one click.",
     "status": "Live"},
    {"name": "Client Management (CRM)",
     "desc": "Full client profiles — contact info, services, payment status, history, internal notes; edit, block, terminate.",
     "status": "Live"},
    {"name": "Applicant Management (Bewerber)",
     "desc": "View employee applications, download submitted documents (CV, passport, driver's license, police clearance), schedule Calendly interviews, approve/reject.",
     "status": "Live"},
    {"name": "Employee Management (Mitarbeiter)",
     "desc": "Roster with document status (not_sent / pending / approved / rejected), resend documents, view earnings, search and filter.",
     "status": "Live"},
    {"name": "Manual Employee Creation",
     "desc": "Add an employee directly without registration form — set details, services offered, availability.",
     "status": "Live"},
    {"name": "Assignments (Einsätze)",
     "desc": "All service assignments — pending, confirmed, completed, cancelled. Manually assign, reassign, cancel; tracks 36h auto-cancel timer.",
     "status": "Live"},
    {"name": "Vacations",
     "desc": "Approve or decline employee and client vacation requests with date-range validation.",
     "status": "Live"},
    {"name": "Finance & Payments (Finanzen)",
     "desc": "All transactions — client charges and employee payouts, status filters, retry logic, manual capture; payment-attempt history.",
     "status": "Live"},
    {"name": "Blog Management",
     "desc": "Create, edit, publish blog articles for the public site (rich-text editor).",
     "status": "Live"},
    {"name": "System Email Broadcast",
     "desc": "Compose and send maintenance notices or announcements to all clients.",
     "status": "Live"},
    {"name": "Feedback Email Trigger",
     "desc": "Manually send a feedback request to a specific client; view submitted ratings.",
     "status": "Live"},
    {"name": "Internal Notes & Tasks",
     "desc": "Searchable notes per user/employee; task list with due dates and completion tracking.",
     "status": "Live"},
    {"name": "Activity Log / Audit Trail",
     "desc": "Records logins, assignment changes, payment attempts, email sends and admin actions with timestamp.",
     "status": "Live"},
    {"name": "Email Template Editor",
     "desc": "UI to view/edit email subject and body templates.",
     "status": "Partial"},
    {"name": "Settings",
     "desc": "System configuration panel.",
     "status": "Partial"},
]

CLIENT_FEATURES = [
    {"name": "Registration & Onboarding",
     "desc": "Multi-step signup: account → personal info → service preferences → health/care needs → payment method. Progress flagged at every step.",
     "status": "Live"},
    {"name": "Dashboard Overview",
     "desc": "Next appointment, assigned caregiver (name, photo, phone), service type, frequency.",
     "status": "Live"},
    {"name": "Appointments / Einsätze",
     "desc": "Calendar and list view of upcoming services.",
     "status": "Live"},
    {"name": "Service History",
     "desc": "Completed appointments with employee, date, service and feedback option.",
     "status": "Live"},
    {"name": "Personal Information",
     "desc": "Edit profile, address, emergency contact, allergies, special requests, household, languages, pets, mobility, health conditions, dietary preferences.",
     "status": "Live"},
    {"name": "Vacation Requests",
     "desc": "Submit, view and track vacation periods; admin approval workflow.",
     "status": "Live"},
    {"name": "Booking & Payment",
     "desc": "Book a service, choose dates/frequency/duration, pay via Stripe (CardElement); confirmation email sent.",
     "status": "Live"},
    {"name": "Saved Payment Methods",
     "desc": "Add/replace card, set default, view past charges and upcoming payments.",
     "status": "Live"},
    {"name": "Cancellation & Refunds",
     "desc": "Cancel booking from dashboard — automatic refund tier: 100% (>7d), 80% (>3d), 50% (<3d).",
     "status": "Live"},
    {"name": "Feedback & Ratings",
     "desc": "1–5 star rating with comment after each service.",
     "status": "Live"},
    {"name": "Password Reset",
     "desc": "Self-service forgot-password link valid 1 hour.",
     "status": "Live"},
    {"name": "Messages",
     "desc": "Communication thread with assigned caregiver / admin.",
     "status": "Partial"},
]

EMPLOYEE_FEATURES = [
    {"name": "Application & Onboarding",
     "desc": "Multi-step employee signup: account → personal info → experience → services offered → languages → availability → required documents (CV, passport, driver's license, police clearance, residence permit/visa).",
     "status": "Live"},
    {"name": "Document Upload",
     "desc": "Submit and re-submit required documents; admin reviews and approves.",
     "status": "Live"},
    {"name": "Interview Scheduling",
     "desc": "Receive Calendly link in invitation email and book an interview slot.",
     "status": "Live"},
    {"name": "Pending Assignments",
     "desc": "View new job proposals — client, service, date, time, duration, location, estimated pay; 36h to accept or reject.",
     "status": "Live"},
    {"name": "Confirmed Assignments",
     "desc": "Active jobs with client contact info, special needs, prep notes and downloadable contract PDF; mark complete after service.",
     "status": "Live"},
    {"name": "Rejected Assignments",
     "desc": "History of declined jobs — system tracks rejection count (warning email at 2+).",
     "status": "Live"},
    {"name": "Vacation Requests",
     "desc": "Submit vacation dates with reason; once approved, system blocks new assignments on those dates.",
     "status": "Live"},
    {"name": "Earnings / Payment Summary",
     "desc": "Hours worked, gross earnings (CHF), pending vs. paid payouts, next payout date, transaction history.",
     "status": "Live"},
    {"name": "Bank Details",
     "desc": "Configure IBAN, BIC, account holder, bank name for payout.",
     "status": "Live"},
    {"name": "Contracts (Rahmenvereinbarung + Einsatz-Arbeitsvertrag)",
     "desc": "Receive framework agreement on approval, plus an individual assignment contract for every confirmed job.",
     "status": "Live"},
    {"name": "Password Reset",
     "desc": "Self-service forgot-password link valid 1 hour.",
     "status": "Live"},
    {"name": "Messages",
     "desc": "Communication thread with clients/admin.",
     "status": "Partial"},
]

PUBLIC_FEATURES = [
    {"name": "Homepage",
     "desc": "Hero, services overview, CTAs to register as client or apply as caregiver, testimonials, blog preview.",
     "status": "Live"},
    {"name": "Services Page",
     "desc": "Detailed service catalogue: daily care, shopping, companionship, medical support, household help — pricing and benefits.",
     "status": "Live"},
    {"name": "Blog",
     "desc": "Public-facing care articles, generated from CMS.",
     "status": "Live"},
    {"name": "About / Contact",
     "desc": "Company info plus contact form (sends inquiry email to admin).",
     "status": "Live"},
    {"name": "Login",
     "desc": "Email/password authentication with role-based redirect (admin / client / employee).",
     "status": "Live"},
]

INTEGRATIONS = [
    ("Stripe", "Payment intents, setup intents, saved payment methods, retries — drives all client payments and refunds."),
    ("Salesforce CRM", "One-way sync: clients to Account, employees to Caregivers__c custom object."),
    ("Mailchimp", "Newsletter subscriptions on registration / opt-in."),
    ("Calendly", "Embedded interview scheduling links in applicant emails."),
    ("Firebase", "Optional auth provider integrated alongside primary password auth."),
    ("Nodemailer (SMTP)", "Powers every transactional email; HTML templates with PDF attachments."),
    ("PDFKit", "Generates Rahmenvereinbarung, Einsatz-Arbeitsvertrag and Abschluss form PDFs."),
    ("Node-Cron", "Scheduled tasks: payment capture, assignment reminders / auto-cancel, capacity emails, unassigned-client notices."),
]

TECH_STACK = [
    ("Framework", "Next.js 14 (full-stack React, API routes)"),
    ("Frontend", "React 18, TailwindCSS, Bootstrap, React-Big-Calendar, React-DatePicker, React-Quill, Chart.js"),
    ("Database", "PostgreSQL (Render) via Prisma 6 ORM"),
    ("Auth", "Password hash (DB) + Firebase optional; role-based access (admin / client / employee)"),
    ("Payments", "Stripe (intents, setup intents, saved methods, webhooks)"),
    ("Email", "Nodemailer over SMTP, HTML templates, PDF attachments"),
    ("PDFs", "PDFKit"),
    ("CRM", "Salesforce via jsforce"),
    ("Marketing", "Mailchimp"),
    ("Scheduling", "Calendly (interview booking)"),
    ("Cron jobs", "Node-Cron (payments capture, assignment reminders, capacity emails)"),
]

# ----------------------------------------------------------------------
# BUILD
# ----------------------------------------------------------------------

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---- Cover ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Prime Home Care")
run.bold = True
run.font.size = Pt(34)
run.font.color.rgb = BRAND

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Platform Documentation")
r.font.size = Pt(20)
r.font.color.rgb = ACCENT

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Email System  •  Admin View  •  Client View  •  Employee View")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GREY

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Version 1.0  •  April 2026")
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_page_break()

# ---- Executive Summary ----
add_heading(doc, "1. Executive Summary", level=1)
add_para(
    doc,
    "Prime Home Care (PHC) is a full-stack elderly home-care management "
    "platform that connects clients with qualified caregivers across "
    "Switzerland. This document describes everything that is fully "
    "operational on the platform today — from the public website and "
    "registration flows through to admin operations, client self-service, "
    "caregiver dispatch and the complete email-notification system that "
    "ties them together.",
)
add_para(
    doc,
    "The platform serves three distinct roles — Administrator, Client and "
    "Employee (Caregiver) — each with its own dashboard, permissions and "
    "workflows. All transactional communication is automated via email, "
    "and key business processes (payment capture, assignment dispatch, "
    "reminders, auto-cancellation) run on scheduled cron jobs.",
)

add_heading(doc, "Platform at a glance", level=2)
add_kv_table(doc, [
    ("Roles supported", "Administrator, Client, Employee"),
    ("Total automated emails", f"{len(EMAILS)} distinct email types"),
    ("Admin features", f"{len([f for f in ADMIN_FEATURES if f['status']=='Live'])} live, {len([f for f in ADMIN_FEATURES if f['status']=='Partial'])} partial"),
    ("Client features", f"{len([f for f in CLIENT_FEATURES if f['status']=='Live'])} live, {len([f for f in CLIENT_FEATURES if f['status']=='Partial'])} partial"),
    ("Employee features", f"{len([f for f in EMPLOYEE_FEATURES if f['status']=='Live'])} live, {len([f for f in EMPLOYEE_FEATURES if f['status']=='Partial'])} partial"),
    ("External integrations", f"{len(INTEGRATIONS)} (Stripe, Salesforce, Mailchimp, Calendly, Firebase, Nodemailer, PDFKit, Node-Cron)"),
    ("Primary language", "German (transactional UI and emails); English available on selected marketing pages"),
])

doc.add_page_break()

# ---- Tech Stack ----
add_heading(doc, "2. Technology Stack", level=1)
add_para(doc, "The platform is built on a modern, production-ready stack:")
add_kv_table(doc, TECH_STACK)

doc.add_page_break()

# ---- Email System ----
add_heading(doc, "3. Email System — Complete Inventory", level=1)
add_para(
    doc,
    "Every transactional email currently sent by the platform is listed "
    "below. All emails are delivered through Nodemailer over SMTP with "
    "branded HTML templates; contracts and framework agreements are "
    "attached as PDFs generated on the fly.",
)
add_emails_table(doc, EMAILS)

add_heading(doc, "Highlight: end-to-end caregiver assignment flow", level=2)
add_bullet(doc, "Admin assigns a caregiver to a confirmed booking.", "1. ")
add_bullet(doc, "Employee receives \"Neuer Einsatz – Bitte bestätigen\".", "2. ")
add_bullet(doc, "If no response in 12h → first reminder email.", "3. ")
add_bullet(doc, "If no response in 24h → final reminder email.", "4. ")
add_bullet(doc, "If still no response in 36h → assignment auto-cancelled, admin notified.", "5. ")
add_bullet(doc, "Once accepted → client receives \"Ihr Betreuer wurde zugewiesen\" and the employee receives the individual assignment contract PDF.", "6. ")

add_heading(doc, "Highlight: client cancellation refund tiers", level=2)
add_bullet(doc, "More than 7 days before service: 100% refund.")
add_bullet(doc, "Between 3 and 7 days before service: 80% refund.")
add_bullet(doc, "Less than 3 days before service: 50% refund.")
add_para(doc, "Each tier produces its own automatically-calculated cancellation confirmation email.", italic=True, color=GREY, size=10)

doc.add_page_break()

# ---- Admin View ----
add_heading(doc, "4. Admin View", level=1)
add_para(
    doc,
    "The Admin Dashboard is the operational cockpit of Prime Home Care. "
    "It gives staff a complete view of the business — pending approvals, "
    "active clients, the caregiver roster, every assignment, finances, "
    "vacations, content and outbound communications.",
)
add_screenshot_placeholder(doc, "Admin Dashboard — Overview")
add_features_table(doc, ADMIN_FEATURES)

add_heading(doc, "Key admin workflows", level=2)
add_bullet(doc, "Approving a new caregiver application: review documents → click Approve → system emails the caregiver with the Rahmenvereinbarung PDF and activates their portal access.")
add_bullet(doc, "Dispatching a service: open the booking → assign a caregiver → system fires the proposal email and starts the 36h timer.")
add_bullet(doc, "Handling a cancellation: cancel from the assignment view → system emails client (with refund tier), employee and admin.")
add_bullet(doc, "Broadcasting maintenance: System Email page → compose → send to all clients in one action.")

doc.add_page_break()

# ---- Client View ----
add_heading(doc, "5. Client View", level=1)
add_para(
    doc,
    "Clients sign up through a guided multi-step form, manage their "
    "personal and household information, book services, pay securely with "
    "Stripe, track every appointment, request vacations and rate the "
    "service afterwards — all from a single dashboard.",
)
add_screenshot_placeholder(doc, "Client Dashboard — Home")
add_features_table(doc, CLIENT_FEATURES)

add_heading(doc, "Client journey", level=2)
add_bullet(doc, "Registration: 4-step form (account, personal info, service preferences, health/care needs) followed by payment method setup.")
add_bullet(doc, "First booking: choose service, dates and frequency → pay via Stripe → receive payment confirmation email.")
add_bullet(doc, "Awaiting caregiver: client receives \"Ihr Betreuer wurde zugewiesen\" once the caregiver is confirmed; if no caregiver is available, daily availability emails are triggered.")
add_bullet(doc, "Day of service: 4-hour reminder email; client meets caregiver as scheduled.")
add_bullet(doc, "After service: feedback request email — 1–5 star rating with comment.")

doc.add_page_break()

# ---- Employee View ----
add_heading(doc, "6. Employee (Caregiver) View", level=1)
add_para(
    doc,
    "Caregivers apply, upload documents, schedule an interview through "
    "Calendly, and once approved gain access to the employee portal where "
    "they receive job proposals, manage assignments, request vacations, "
    "track earnings and download their contracts.",
)
add_screenshot_placeholder(doc, "Employee Dashboard — Home")
add_features_table(doc, EMPLOYEE_FEATURES)

add_heading(doc, "Caregiver journey", level=2)
add_bullet(doc, "Application: complete the multi-step form and upload required documents (CV, passport, driver's license, police clearance, residence permit/visa).")
add_bullet(doc, "Interview: receive Calendly invitation email and book a slot.")
add_bullet(doc, "Approval: admin approves → caregiver receives Rahmenvereinbarung PDF and portal access.")
add_bullet(doc, "Receiving jobs: each new assignment proposal lands by email and in Pending Assignments — accept or reject within 36 hours.")
add_bullet(doc, "Working a job: confirmed assignment shows full client info and includes a downloadable Einsatz-Arbeitsvertrag PDF.")
add_bullet(doc, "Getting paid: earnings page shows hours, gross pay (CHF), pending vs. paid payouts and the next payout date.")

doc.add_page_break()

# ---- Public Pages ----
add_heading(doc, "7. Public Pages", level=1)
add_features_table(doc, PUBLIC_FEATURES)

# ---- Integrations ----
add_heading(doc, "8. External Integrations", level=1)
add_kv_table(doc, INTEGRATIONS)

# ---- Scheduled Jobs ----
add_heading(doc, "9. Scheduled Background Jobs", level=1)
add_para(doc, "Recurring tasks that keep the platform running without human intervention:")
add_bullet(doc, "Payment capture cron — every 10 minutes; retries failed Stripe charges and emails clients on payment failure.", "• ")
add_bullet(doc, "Assignment reminder cron — sends 12h and 24h reminders to caregivers and auto-cancels at 36h.", "• ")
add_bullet(doc, "Daily capacity cron — emails clients when no caregiver capacity is available for a booked period.", "• ")
add_bullet(doc, "Unassigned-client cron — notifies clients still waiting for a caregiver assignment.", "• ")
add_bullet(doc, "Service reminder cron — 48h profile-completion reminder and 4h pre-service reminder.", "• ")
add_bullet(doc, "Feedback request cron — sends post-service rating requests.", "• ")

# ---- Status Legend ----
add_heading(doc, "10. Status Legend", level=1)
p = doc.add_paragraph()
r = p.add_run("Live")
r.bold = True
r.font.color.rgb = GREEN
p.add_run(" — feature is fully implemented, tested and used in day-to-day operations.")
p = doc.add_paragraph()
r = p.add_run("Partial")
r.bold = True
r.font.color.rgb = ORANGE
p.add_run(" — UI is in place but some backend wiring or polish is still in progress.")

# ---- Footer note ----
doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run(
    "Prime Home Care AG  •  Birkenstrasse 49, CH-6343  •  phc.ch"
)
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

OUT = r"c:\Users\Lenovo\Desktop\PHC-final\PHC_Platform_Documentation.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
