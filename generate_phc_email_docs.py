"""
Generates PHC_Email_System_Documentation.docx — a complete English Word document
covering every email the platform sends: who receives it, when it fires,
which subject and body it uses, what it attaches and which cron drives it.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x0B, 0x4F, 0x6C)
ACCENT = RGBColor(0x1A, 0x76, 0xD2)
GREY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
ORANGE = RGBColor(0xE6, 0x5A, 0x00)
RED = RGBColor(0xC0, 0x39, 0x2B)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=BRAND):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 0:
        run.font.size = Pt(26)
    elif level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_kv_table(doc, rows, col_widths=(Cm(5), Cm(11.5))):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.text = ""
        c2.text = ""
        r1 = c1.paragraphs[0].add_run(k)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(v)
        r2.font.size = Pt(10)
        c1.width = col_widths[0]
        c2.width = col_widths[1]
        set_cell_bg(c1, "EAF2F8")
    return table


def styled_table(doc, headers, rows, header_bg="0B4F6C", body_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], header_bg)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(body_size)
    return table


def email_card(doc, n, title, fields):
    """Render a single email as a heading + key/value mini-table."""
    h = doc.add_paragraph()
    rn = h.add_run(f"{n}. {title}")
    rn.bold = True
    rn.font.size = Pt(12)
    rn.font.color.rgb = BRAND
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    add_kv_table(doc, fields, col_widths=(Cm(3.5), Cm(13)))


# ----------------------------------------------------------------------
# DATA — the complete catalogue of every email send-site
# ----------------------------------------------------------------------

EMAILS = [
    {
        "title": "Client Welcome — \"Willkommen bei Prime Home Care – Ihr Zugang ist aktiv\"",
        "trigger": "A new client completes registration; the system creates the client portal account and sends a welcome with quick-start info.",
        "when": "Immediately after successful client registration (POST /api/register with role=client).",
        "to": "Client (email captured at registration)",
        "subject": "Willkommen bei Prime Home Care – Ihr Zugang ist aktiv",
        "body": "Welcome message, list of what they can do in the portal (manage bookings, message us, view appointments), contact info.",
        "attach": "None",
        "where": "src/pages/api/register.js (client branch)",
    },
    {
        "title": "Employee Welcome — \"Willkommen bei Prime Home Care AG – Passwort erstellen\"",
        "trigger": "Admin or system creates a new employee account with role=employee; system generates a password-setup link valid 7 days.",
        "when": "On POST /api/register with role=employee.",
        "to": "Employee",
        "subject": "Willkommen bei Prime Home Care AG – Passwort erstellen",
        "body": "Welcome to the team, button to set initial password (7-day token), branding and contact footer.",
        "attach": "None",
        "where": "src/pages/api/register.js — sendSetupEmail()",
    },
    {
        "title": "Application Received — \"Ihre Bewerbung bei Prime Home Care AG\"",
        "trigger": "Caregiver submits the public job-application form (with CV) on the landing page.",
        "when": "On POST /api/sendEmail (form submission).",
        "to": "Admin (landingpage@phc.ch) + CC admin@phc.ch + jobs@phc.ch",
        "subject": "Formular: Jobs Landing Page {applicant name}",
        "body": "Applicant details (Vorname, Nachname, email, region, message) formatted as an HTML table for the recruiting team.",
        "attach": "CV file if uploaded (5MB limit)",
        "where": "src/pages/api/sendEmail.js",
    },
    {
        "title": "Generic Job Inquiry — \"Formular: Jobs Landing Page\"",
        "trigger": "Visitor uses the simpler \"contact us about jobs\" form (no CV) on the landing page.",
        "when": "On POST /api/sendMail.",
        "to": "Admin (landingpage@phc.ch) + CC edita.latifi@the-eksperts.com + jobs@phc.ch",
        "subject": "Formular: Jobs Landing Page {email}",
        "body": "Plain HTML showing the email and the question/inquiry text.",
        "attach": "None",
        "where": "src/pages/api/sendMail.js",
    },
    {
        "title": "Interview Invitation (Calendly) — \"Terminvereinbarung für ein persönliches Kennenlernen\"",
        "trigger": "Admin invites an applicant to interview from the Bewerber view; system pulls a stored email template and embeds the Calendly link.",
        "when": "On POST /api/invite-employee. Marks applicant.invited=true and stores inviteSentAt.",
        "to": "Applicant (Employee table)",
        "subject": "From DB template `interviewInvite.subject`",
        "body": "Templated body with placeholders {{firstName}} and {{calendlyLink}} (https://calendly.com/primehomecare).",
        "attach": "None",
        "where": "src/pages/api/invite-employee.js",
    },
    {
        "title": "Interview Invitation (direct) — \"Vereinbaren Sie Ihr Interview\"",
        "trigger": "Alternative direct-Calendly invitation utility.",
        "when": "Helper called from admin flows when the templated path is bypassed.",
        "to": "Applicant",
        "subject": "Vereinbaren Sie Ihr Interview",
        "body": "Short hard-coded HTML inviting the candidate to book a slot on Calendly.",
        "attach": "None",
        "where": "src/lib/sendCalendlyInvite.js",
    },
    {
        "title": "Interview Booking Reminder — \"Erinnerung: Termin für Kennenlernen noch nicht vereinbart\"",
        "trigger": "Applicant has invited=true and inviteSentAt is older than 3 days but they have not booked a Calendly slot.",
        "when": "Background job — runs daily via the scheduler; only applicants with status=pending qualify.",
        "to": "Applicant",
        "subject": "Erinnerung: Termin für Kennenlernen noch nicht vereinbart",
        "body": "Friendly reminder to use the Calendly link from the original invitation.",
        "attach": "None",
        "where": "src/lib/sendReminderForInviteInterview.js",
    },
    {
        "title": "Application Approved — \"Willkommen im Prime Home Care Team – Ihr Zugang ist aktiviert\"",
        "trigger": "Admin clicks Approve on an applicant; status changes to approved.",
        "when": "On POST /api/approve-employee. Idempotent — only fires if !previouslyApproved.",
        "to": "Employee",
        "subject": "Willkommen im Prime Home Care Team – Ihr Zugang ist aktiviert",
        "body": "Approval confirmation, login URL, hints to start; references the attached framework agreement.",
        "attach": "Rahmenvereinbarung PDF (createRahmenvereinbarungPdf in src/lib/mailer.js)",
        "where": "src/pages/api/approve-employee.js → src/lib/mailer.js sendApprovalEmail()",
    },
    {
        "title": "Application Rejected — DB template `rejectionEmail`",
        "trigger": "Admin rejects an applicant from the Bewerber view; status changes to rejected.",
        "when": "On POST /api/reject-employee.",
        "to": "Applicant",
        "subject": "From DB template `rejectionEmail.subject`",
        "body": "Polite rejection message; thanks for the application; placeholders {{firstName}}.",
        "attach": "None",
        "where": "src/pages/api/reject-employee.js",
    },
    {
        "title": "New Assignment Proposal — DB template `assignmentNotification`",
        "trigger": "Admin assigns a caregiver to a confirmed booking from the Einsätze or Clients page.",
        "when": "On POST /api/assign or /api/admin/assign-client. Starts the 36h confirmation clock.",
        "to": "Employee assigned",
        "subject": "From DB template `assignmentNotification.subject`",
        "body": "Job details (client, service, date, time, duration, location, estimated pay), \"Accept / Reject\" CTAs to the dashboard. Placeholders {{firstName}} etc.",
        "attach": "None",
        "where": "src/pages/api/assign.js + src/pages/api/admin/assign-client.js",
    },
    {
        "title": "Assignment Reminder 12h — \"Reminder: Bitte Job akzeptieren\"",
        "trigger": "Cron — assignment.confirmationStatus=pending, reminderSentCount=0 and assignment was created ≥ 12h ago.",
        "when": "Every minute by the cron job; first reminder fires once 12h elapses with no response.",
        "to": "Employee",
        "subject": "Reminder: Bitte Job akzeptieren",
        "body": "Reminder to open the dashboard and confirm or reject the proposal; bumps reminderSentCount to 1.",
        "attach": "None",
        "where": "src/pages/api/cron/assignment-cron.js (runAssignmentReminders)",
    },
    {
        "title": "Assignment Reminder 24h (final) — \"Letzte Erinnerung: Bitte Job akzeptieren\"",
        "trigger": "Cron — confirmationStatus=pending, reminderSentCount=1, createdAt ≥ 24h.",
        "when": "Every minute by the cron job; second/final reminder fires once 24h has passed.",
        "to": "Employee",
        "subject": "Letzte Erinnerung: Bitte Job akzeptieren",
        "body": "Final reminder before auto-cancellation; bumps reminderSentCount to 2.",
        "attach": "None",
        "where": "src/pages/api/cron/assignment-cron.js (runAssignmentReminders)",
    },
    {
        "title": "Assignment Auto-Cancelled — \"Einsatz automatisch storniert\"",
        "trigger": "Cron — confirmationStatus=pending, reminderSentCount ≥ 2, createdAt ≥ 36h. System auto-cancels and clears employeeId on the schedule.",
        "when": "Every minute by the cron job; admin notification is the only email at this stage (the schedule frees up automatically).",
        "to": "All admin users (User where role=admin)",
        "subject": "Einsatz automatisch storniert",
        "body": "Admin notice naming the employee and the cancelled assignment so a new caregiver can be assigned.",
        "attach": "None",
        "where": "src/pages/api/cron/assignment-cron.js",
    },
    {
        "title": "Caregiver Assigned — \"Ihre Buchung wurde erfolgreich bestätigt\"",
        "trigger": "Employee accepts an assignment.",
        "when": "On POST /api/employee/confirm-assignment with action=confirmed (and not already confirmed).",
        "to": "Client",
        "subject": "Ihre Buchung wurde erfolgreich bestätigt",
        "body": "Confirms the booking, gives the caregiver's name, phone, service name and first service date.",
        "attach": "None",
        "where": "src/lib/mailer.js → sendAssignmentAcceptedEmail()",
    },
    {
        "title": "Assignment Contract — Einsatz-Arbeitsvertrag PDF",
        "trigger": "Employee accepts an assignment (same trigger as above).",
        "when": "Same call to /api/employee/confirm-assignment — the system also generates and emails the individual contract.",
        "to": "Employee",
        "subject": "Einsatz-Arbeitsvertrag",
        "body": "Cover note attaching the per-assignment work contract PDF (employee, client, service, date/time, duration, pay, T&Cs).",
        "attach": "Einsatz-Arbeitsvertrag PDF (createContractPdf in src/lib/emailHelpers.js)",
        "where": "src/lib/emailHelpers.js → sendAssignmentContractEmail()",
    },
    {
        "title": "Rejection Warning — DB template `rejectionWarning`",
        "trigger": "Employee rejects an assignment AND has rejected ≥ 3 assignments AND no warning has been sent before.",
        "when": "On POST /api/employee/confirm-assignment with action=rejected and the threshold met.",
        "to": "Employee + CC admin@phc.ch",
        "subject": "From DB template `rejectionWarning.subject`",
        "body": "Templated written warning explaining the rejection pattern and possible consequences.",
        "attach": "None",
        "where": "src/pages/api/employee/confirm-assignment.js",
    },
    {
        "title": "Booking Capacity Notice — \"Information zur Verfügbarkeit Ihrer Buchung\"",
        "trigger": "Cron — client created a booking 48h ago and still has no caregiver assigned.",
        "when": "POST /api/cron (CRON_SECRET-protected) calls runUnassignedClientEmails().",
        "to": "Client",
        "subject": "Information zur Verfügbarkeit Ihrer Buchung",
        "body": "Polite notice that no caregiver capacity is currently available for the requested period; offers to keep trying.",
        "attach": "None",
        "where": "src/scripts/unassigned-client-cron.js",
    },
    {
        "title": "Two-day Service Reminder — \"In 2 Tagen beginnt Ihre Betreuung\"",
        "trigger": "Background job — schedule.date is exactly 2 days from today.",
        "when": "Daily run of notifySchedulesInTwoDays().",
        "to": "Client",
        "subject": "In 2 Tagen beginnt Ihre Betreuung.",
        "body": "Reminder of the upcoming service date with caregiver contact info and service details.",
        "attach": "None",
        "where": "src/lib/sendRemindersEmailClient.js",
    },
    {
        "title": "48h Profile-Completion Reminder — \"Bitte fehlende Informationen in der Checkliste ergänzen\"",
        "trigger": "Reminder row of type=48h_reminder with sent=false and scheduledAt ≤ now.",
        "when": "Background scheduler running sendPendingReminders().",
        "to": "Client",
        "subject": "Bitte fehlende Informationen in der Checkliste ergänzen",
        "body": "Asks the client to finish the household / care checklist before the caregiver arrives.",
        "attach": "None",
        "where": "src/lib/sendReminders.js",
    },
    {
        "title": "4h Service Reminder — \"Ihre Betreuung wartet auf Sie\"",
        "trigger": "Reminder row of type=4h_reminder with sent=false and scheduledAt ≤ now.",
        "when": "Background scheduler running sendPendingReminders().",
        "to": "Client",
        "subject": "Ihre Betreuung wartet auf Sie",
        "body": "Last-minute reminder; appointment in 4 hours.",
        "attach": "None",
        "where": "src/lib/sendReminders.js",
    },
    {
        "title": "Payment Confirmation — \"Zahlungsbestätigung / Rechnung zu Ihrer Buchung\"",
        "trigger": "Stripe charge succeeds during cron-driven payment capture (assignment firstDate ≤ 24h away, paymentCaptured=false).",
        "when": "POST /api/cron → chargeUnpaidUsers() → capturePendingPayments().",
        "to": "Client",
        "subject": "Zahlungsbestätigung / Rechnung zu Ihrer Buchung",
        "body": "Receipt / invoice — service, dates, amount in CHF, transaction reference.",
        "attach": "None",
        "where": "src/scripts/capturePayments.js — sendPaymentConfirmationEmail()",
    },
    {
        "title": "Payment Failure Alert — \"⚠️ Payment failed 3 times – {Name}\"",
        "trigger": "Retry cron — Stripe charge has failed and the assignment now has attempts ≥ 3.",
        "when": "POST /api/finances/retry-payments (CRON_SECRET-protected).",
        "to": "Admin (process.env.ADMIN_NOTIFICATION_EMAIL)",
        "subject": "⚠️ Payment failed 3 times – {firstName} {lastName}",
        "body": "Names the client, the amount and the Stripe error so the admin can intervene manually.",
        "attach": "None",
        "where": "src/pages/api/finances/retry-payments.js",
    },
    {
        "title": "Payment Change Request — \"Zahlungsänderung angefragt\"",
        "trigger": "Client submits the \"change my payment method\" form from the dashboard.",
        "when": "On POST /api/request-payment-change.",
        "to": "Admin (admin@phc.ch)",
        "subject": "Zahlungsänderung angefragt",
        "body": "Identifies the client (email + name) so admin can follow up and update Stripe details.",
        "attach": "None",
        "where": "src/pages/api/request-payment-change.js",
    },
    {
        "title": "Termination — Ordinary Notice — \"Bestätigung Ihrer Kündigung bei Prime Home Care AG\"",
        "trigger": "Client cancels their account with a notice period (Abschluss form, immediate=false).",
        "when": "Triggered from the client termination flow.",
        "to": "Client",
        "subject": "Bestätigung Ihrer Kündigung bei Prime Home Care AG",
        "body": "Confirms the cancellation, the end date and the final billing window.",
        "attach": "None",
        "where": "src/lib/sendTerminationEmail.js (immediate=false branch)",
    },
    {
        "title": "Termination — Immediate — \"Bestätigung Ihrer fristlosen Kündigung\"",
        "trigger": "Client requests immediate (fristlose) termination.",
        "when": "Triggered from the client termination flow with immediate=true.",
        "to": "Client",
        "subject": "Bestätigung Ihrer fristlosen Kündigung bei Prime Home Care AG",
        "body": "Confirms immediate deactivation and references the CHF 300 administrative fee (Aufwandsentschädigung).",
        "attach": "None",
        "where": "src/lib/sendTerminationEmail.js (immediate=true branch)",
    },
    {
        "title": "Password Reset — \"Passwort zurücksetzen\"",
        "trigger": "User clicks Forgot Password (works for both client and employee email lookup).",
        "when": "On POST /api/forgot-password. Always returns 200 to prevent email enumeration.",
        "to": "Client or Employee",
        "subject": "Passwort zurücksetzen",
        "body": "Branded HTML with a 1-hour reset link, company address footer and links to AVB / Nutzungsbedingungen.",
        "attach": "None",
        "where": "src/pages/api/forgot-password.js",
    },
    {
        "title": "Caregiver Feedback Request — DB template `feedbackEmail`",
        "trigger": "Admin clicks Send Feedback Request for a specific caregiver from the admin tools.",
        "when": "On POST /api/admin/send-feedback-email — broadcast to all clients in parallel via Promise.all.",
        "to": "All clients (role=client)",
        "subject": "From DB template `feedbackEmail.subject`",
        "body": "Templated request to rate a caregiver. Placeholders: {{firstName}}, {{lastName}}, {{caregiverName}}, {{feedbackLink}}.",
        "attach": "None",
        "where": "src/pages/api/admin/send-feedback-email.js",
    },
    {
        "title": "Newsletter Subscription (Mailchimp)",
        "trigger": "Visitor opts into the newsletter on the public site.",
        "when": "On POST /api/subscribe — POSTs to Mailchimp /3.0/lists/{LIST_ID}/members.",
        "to": "Subscriber (Mailchimp delivers the welcome / opt-in email itself)",
        "subject": "Mailchimp default — managed in the Mailchimp audience settings",
        "body": "Mailchimp double-opt-in / welcome automation; nothing rendered server-side by PHC.",
        "attach": "—",
        "where": "src/pages/api/subscribe.js",
    },
]


CRONS = [
    [
        "assignment-reminder-cron.js",
        "* * * * * (every minute)",
        "12h reminder, 24h final reminder, 36h auto-cancel + admin notice",
        "node-cron schedule registered inside the route",
    ],
    [
        "/api/cron",
        "External HTTP (CRON_SECRET)",
        "Payment capture + payment confirmation; unassigned-client capacity notices",
        "Caller (e.g. Render cron, Vercel cron) hits the endpoint",
    ],
    [
        "/api/finances/retry-payments",
        "External HTTP (CRON_SECRET)",
        "Stripe retry, then admin alert at attempts ≥ 3",
        "Caller schedules the retry cadence (typically every 48h)",
    ],
    [
        "Background scheduler — sendRemindersEmailClient",
        "Daily",
        "2-day pre-service client reminder",
        "Runs notifySchedulesInTwoDays() against tomorrow+1 schedules",
    ],
    [
        "Background scheduler — sendReminders",
        "Frequent (per-minute / per-hour)",
        "48h checklist reminder, 4h pre-service reminder",
        "Reads pending Reminder rows where sent=false and scheduledAt ≤ now()",
    ],
    [
        "Background scheduler — sendReminderForInviteInterview",
        "Daily",
        "Interview-booking reminder if Calendly slot not booked within 3 days",
        "Filters Employees where invited=true, status=pending, inviteSentAt ≥ 72h ago",
    ],
]


JOURNEYS = [
    (
        "Caregiver application → approval",
        [
            "Application Received (admin notification with CV)",
            "Interview Invitation (Calendly link)",
            "Interview Booking Reminder (if no booking within 3 days)",
            "Approval email + Rahmenvereinbarung PDF, OR Rejection email",
            "Welcome / password setup link if account is created via /api/register",
        ],
    ),
    (
        "Client booking → service → feedback",
        [
            "Welcome email at registration",
            "Payment Confirmation when Stripe captures the charge",
            "Caregiver Assigned email once an employee accepts",
            "48h profile-completion reminder",
            "2-day pre-service reminder",
            "4h pre-service reminder",
            "Capacity Notice if no caregiver is found within 48h",
            "Caregiver Feedback Request (admin-initiated, post-service)",
        ],
    ),
    (
        "Assignment proposal lifecycle (caregiver side)",
        [
            "T+0h: New Assignment Proposal email to employee",
            "T+12h: Reminder email if still pending",
            "T+24h: Final reminder if still pending",
            "T+36h: Auto-cancel — admin notification email; schedule freed",
            "On Accept: Einsatz-Arbeitsvertrag PDF to employee + Caregiver-Assigned email to client",
            "On Reject (3rd+ time): Rejection Warning email to employee, CC admin",
        ],
    ),
    (
        "Payment lifecycle",
        [
            "Stripe success → Payment Confirmation to client",
            "Stripe failure → silent retry (attempts ++ , next retry in 48h)",
            "After 3 failed attempts → Payment Failure Alert to admin",
            "Client requests change of method → Payment Change Request to admin",
        ],
    ),
    (
        "Termination",
        [
            "Ordinary notice → Termination confirmation (with end date)",
            "Immediate (fristlos) → Termination confirmation referencing CHF 300 fee",
        ],
    ),
    (
        "Password reset",
        [
            "Forgot Password form → 1-hour reset link emailed (works for client or employee email)",
            "Always returns 200 server-side — email enumeration is prevented",
        ],
    ),
]


# ----------------------------------------------------------------------
# BUILD
# ----------------------------------------------------------------------

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---- Cover ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Prime Home Care")
run.bold = True
run.font.size = Pt(34)
run.font.color.rgb = BRAND

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Email System Documentation")
r.font.size = Pt(20)
r.font.color.rgb = ACCENT

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run(
    "Every transactional email the platform sends — who, when, why, with what."
)
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GREY

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Version 1.0  •  April 2026")
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_page_break()

# ---- 1. Overview ----
add_heading(doc, "1. Overview", level=1)
add_para(
    doc,
    "Prime Home Care delivers all customer, caregiver and admin communication "
    "by email. There is no SMS, push or paper channel. This document describes "
    "every email the platform produces — what triggers it, who receives it, "
    "what it contains, whether it carries an attachment, and where in the "
    "codebase the send is implemented.",
)
add_para(
    doc,
    f"The platform sends {len(EMAILS)} distinct email types, "
    "driven by a mix of direct user actions (e.g. registration, booking, "
    "Forgot Password) and scheduled background jobs (assignment reminders, "
    "service reminders, payment capture, capacity notices, etc.).",
)

add_heading(doc, "How an email is sent (high level)", level=2)
add_bullet(doc, "Nodemailer over SMTP is the only delivery channel (Mailchimp is used only for the public newsletter).")
add_bullet(doc, "All transactional emails are sent server-side from /api routes or scheduled jobs — never from the browser.")
add_bullet(doc, "Two kinds of templates are used: hard-coded HTML strings inside the codebase, and database-stored templates (table emailTemplate) the admin can edit.")
add_bullet(doc, "Two kinds of attachments are generated on the fly with PDFKit: the Rahmenvereinbarung (framework agreement) and the Einsatz-Arbeitsvertrag (per-assignment contract).")
add_bullet(doc, "Branding: the FROM address is \"Prime Home Care AG\" <SMTP_USER>; emails carry the company footer (Birkenstrasse 49, CH-6343), AVB and Nutzungsbedingungen links.")
add_bullet(doc, "Language: all transactional emails are in German.")

add_heading(doc, "FROM addresses and admin recipients", level=2)
add_kv_table(doc, [
    ("Default FROM", "\"Prime Home Care AG\" <SMTP_USER>"),
    ("Generic admin notifications", "admin@phc.ch"),
    ("Recruitment / job inbound", "jobs@phc.ch"),
    ("Landing-page form submissions", "landingpage@phc.ch"),
    ("Critical payment alerts", "process.env.ADMIN_NOTIFICATION_EMAIL"),
    ("Internal / team CC (job forms)", "edita.latifi@the-eksperts.com"),
])

doc.add_page_break()

# ---- 2. Complete email catalogue ----
add_heading(doc, "2. Complete Email Catalogue", level=1)
add_para(
    doc,
    "Each card below describes one email the system sends. \"When\" is the exact "
    "trigger condition; \"Sent to\" is the recipient role; \"Where in code\" "
    "points to the send-site.",
    italic=True, color=GREY, size=10,
)

for idx, e in enumerate(EMAILS, 1):
    email_card(doc, idx, e["title"], [
        ("Trigger", e["trigger"]),
        ("When it fires", e["when"]),
        ("Sent to", e["to"]),
        ("Subject", e["subject"]),
        ("Body summary", e["body"]),
        ("Attachment", e["attach"]),
        ("Where in code", e["where"]),
    ])

doc.add_page_break()

# ---- 3. Lifecycle journeys ----
add_heading(doc, "3. Lifecycle Journeys", level=1)
add_para(
    doc,
    "Same emails as above, but grouped by the user journey that triggers them. "
    "This is the easiest way to verify end-to-end coverage in a stakeholder "
    "review.",
)
for journey, steps in JOURNEYS:
    add_heading(doc, journey, level=2)
    for s in steps:
        add_bullet(doc, s)

doc.add_page_break()

# ---- 4. Cron / scheduled jobs ----
add_heading(doc, "4. Scheduled Jobs Driving Emails", level=1)
add_para(
    doc,
    "These background processes are responsible for every email that does not "
    "originate from a direct user click. The first job is registered with "
    "node-cron inside the API route; the others are HTTP endpoints invoked by "
    "an external scheduler (e.g. Render cron) on the cadence shown.",
)
styled_table(
    doc,
    headers=["Job / route", "Cadence", "Emails it can fire", "How it is triggered"],
    rows=CRONS,
    body_size=9,
)

# ---- 5. DB-stored templates ----
add_heading(doc, "5. Database Email Templates", level=1)
add_para(
    doc,
    "Six emails are not hard-coded — they live in the EmailTemplate table and "
    "the admin can edit subject and body via the admin UI. Placeholders are "
    "replaced at send time with the per-recipient values.",
)
styled_table(
    doc,
    headers=["Template name", "Used for", "Common placeholders"],
    rows=[
        ["assignmentNotification", "New job proposal to caregiver", "{{firstName}}, {{client}}, {{date}}, {{service}}"],
        ["interviewInvite", "Calendly interview invitation", "{{firstName}}, {{calendlyLink}}"],
        ["rejectionEmail", "Application rejection", "{{firstName}}"],
        ["rejectionWarning", "Caregiver rejected ≥3 jobs", "{{firstName}}, {{count}}"],
        ["feedbackEmail", "Post-service feedback request to all clients", "{{firstName}}, {{lastName}}, {{caregiverName}}, {{feedbackLink}}"],
    ],
    body_size=9,
)

# ---- 6. PDF attachments ----
add_heading(doc, "6. PDF Attachments", level=1)
add_kv_table(doc, [
    ("Rahmenvereinbarung",
     "Framework employment agreement — generated when the admin approves an "
     "applicant; attached to the approval email. Source: createRahmenvereinbarungPdf() "
     "in src/lib/mailer.js."),
    ("Einsatz-Arbeitsvertrag",
     "Per-assignment work contract — generated when a caregiver accepts a "
     "specific job; attached to the assignment-contract email. Source: "
     "createContractPdf() in src/lib/emailHelpers.js."),
    ("Job application CV",
     "Forwarded as-is from the public landing-page form (5 MB upload limit) "
     "to the recruiting inbox."),
])

# ---- 7. Status by audience ----
add_heading(doc, "7. Emails by Audience", level=1)

add_heading(doc, "Client receives", level=2)
client_emails = [
    "Client Welcome",
    "Caregiver Assigned",
    "Capacity Notice (no caregiver yet)",
    "2-day, 48h, 4h pre-service reminders",
    "Payment Confirmation",
    "Termination confirmation (ordinary or immediate)",
    "Password Reset",
    "Caregiver Feedback Request",
]
for e in client_emails:
    add_bullet(doc, e)

add_heading(doc, "Employee / caregiver receives", level=2)
employee_emails = [
    "Application Received confirmation",
    "Interview Invitation (Calendly)",
    "Interview Booking Reminder",
    "Approval email with Rahmenvereinbarung PDF",
    "Rejection email",
    "New Assignment Proposal",
    "12h Reminder, 24h Final Reminder",
    "Einsatz-Arbeitsvertrag PDF on acceptance",
    "Rejection Warning (3+ rejections)",
    "Welcome / password setup link",
    "Password Reset",
]
for e in employee_emails:
    add_bullet(doc, e)

add_heading(doc, "Admin receives", level=2)
admin_emails = [
    "Job application form submissions (with and without CV)",
    "Auto-cancel notification when an assignment lapses at 36h",
    "Payment Failure Alert after 3 failed Stripe attempts",
    "Payment Change Request from a client",
    "Rejection-warning copy when a caregiver is warned",
]
for e in admin_emails:
    add_bullet(doc, e)

# ---- Footer ----
doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run(
    "Prime Home Care AG  •  Birkenstrasse 49, CH-6343  •  phc.ch"
)
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

OUT = r"c:\Users\Lenovo\Desktop\PHC-final\PHC_Email_System_Documentation.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
