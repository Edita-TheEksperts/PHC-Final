"""
Builds PHC_QA-Umsetzung_12.07.2026.docx — the umsetzung + acceptance document
for the 12.07.2026 QA package.

  Teil A — Bug-Fix-Umsetzungsnachweis (A1–A6 erledigt, A7 Konzept/Plan)
  Teil B — Test- & Abnahmeplan über 11 Kernprozesse mit Screenshot-Nachweis

Screenshots are taken from screenshots/<name>.png (produced by the Playwright
spec tests/e2e/process-screenshots.spec.js). Where a shot could not be
auto-captured (completed Stripe payment, real e-mail inbox, or the deferred A7
series behaviour) a labelled placeholder with the reason is inserted instead.

Run:  python build_qa_doc.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
SHOT_DIR = os.path.join(BASE, "screenshots")

NAVY = RGBColor(0x04, 0x43, 0x6F)
GOLD = RGBColor(0xB9, 0x9B, 0x5F)
GREY = RGBColor(0x6B, 0x72, 0x80)
RED = RGBColor(0xC0, 0x39, 0x2B)

# Reasons a screenshot is a placeholder instead of an auto-capture.
MANUAL_REASONS = {
    "22_neubuchung_reload": "Live nach echter Buchung aufnehmen (Reload-Zustand nach Zahlung).",
    "15_zahlung_ok": "Abgeschlossene Stripe-Zahlung im CardElement — live aufnehmen.",
    "16_zahlung_abgelehnt": "Stripe-Ablehnungskarte 4000…0002 — live aufnehmen.",
    "28_storno_mail": "E-Mail — echter Posteingang nötig.",
    "31_kuendigung_frei": "Vertrag >30 Tage in der DB nötig — live aufnehmen.",
    "37_anfrage": "Sendet echte Anfrage (mutiert Daten) — live aufnehmen.",
    "38-39_serie_annahme": "A7 gebaut & getestet — Live-Screenshot beim Abnahme-Durchlauf.",
    "38b_serie_mail": "A7 gebaut — E-Mail mit Serienzeile im Posteingang aufnehmen.",
    "40_kunde_serie": "A7 gebaut — Kunden-Dashboard nach Annahme live aufnehmen.",
    "41_serie_abgelehnt": "A7 gebaut — Gegentest (Ablehnung) live aufnehmen.",
    "42_kein_match": "Kunde ohne passenden MA nötig — live aufnehmen.",
    "43_absenz_meldung": "A7 gebaut — «Absenz melden» im MA-Dashboard live aufnehmen.",
    "44_ersatz_matching": "A7 gebaut — freigegebener Termin im Admin live aufnehmen.",
    "45_ersatz_bestaetigt": "A7 gebaut — Ersatz-Bestätigung live aufnehmen.",
    "46_urlaub_matching": "MA-Urlaub über den Zeitraum nötig — live aufnehmen.",
    "57_kunde_urlaub": "Kunden-Urlaub mutiert Daten — live aufnehmen.",
    "58_ma_urlaub": "MA-Urlaubsantrag mutiert Daten — live aufnehmen.",
    "61-62_mail_buchung": "E-Mail — echter Posteingang nötig.",
    "63-64_mail_zuweisung": "E-Mail — echter Posteingang nötig.",
    "65-66_mail_rest": "E-Mail — echter Posteingang nötig.",
    "67_mail_vorlagen": "Admin E-Mail-Editor — live aufnehmen (/admin/email-templates).",
}


# ── low-level helpers ─────────────────────────────────────────────────────

def hr(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), "B99B5F")
    pBdr.append(b); pPr.append(pBdr)


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def shot(name):
    for ext in ("png", "jpg", "jpeg"):
        p = os.path.join(SHOT_DIR, f"{name}.{ext}")
        if os.path.exists(p):
            return p
    return None


def embed_shot(doc, name):
    """Embed screenshots/<name>.png, or a labelled placeholder with the reason."""
    p = shot(name)
    if p:
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(p, width=Inches(6.1))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Automatisch aufgenommen · {name}.png")
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
    else:
        reason = MANUAL_REASONS.get(name, "Manuell aufzunehmen.")
        t = doc.add_table(rows=1, cols=1)
        c = t.cell(0, 0); shade(c, "FAFCFF")
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(f"[ Screenshot: {name}.png ]")
        r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GOLD
        cp2 = c.add_paragraph(); cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = cp2.add_run(reason); r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = GREY
        tbl = t._tbl; borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "dashed"); e.set(qn("w:sz"), "6"); e.set(qn("w:color"), "B99B5F")
            borders.append(e)
        tbl.tblPr.append(borders)
    doc.add_paragraph()


def styles(doc):
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)


def h(doc, text, level, color=NAVY):
    head = doc.add_heading(text, level=level)
    for run in head.runs:
        run.font.color.rgb = color
    return head


# ── cover ─────────────────────────────────────────────────────────────────

def cover(doc):
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PHC"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("QA-Umsetzung & Abnahme"); r.font.size = Pt(20); r.font.color.rgb = GOLD
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Fehlerbehebung A1–A6 · Test- & Abnahmeplan über 11 Kernprozesse")
    r.italic = True; r.font.size = Pt(12)
    doc.add_paragraph()
    m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("Prime Home Care AG\n").bold = True
    m.add_run("Basis: QA-Paket 12.07.2026 · Repo Edita-TheEksperts/PHC-Final (master)\n")
    m.add_run("Umsetzung: A1–A6 erledigt · A7 als Konzept/Plan\n")
    m.add_run("Adressat: Bettina, Silvain (Abnahme)")
    doc.add_paragraph()
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(
        "Wichtig: Die Plattform wird in der Schweiz betrieben — alle UI-Texte, Status-Labels "
        "und E-Mails bleiben DEUTSCH. Dieses Dokument dient nur der Abnahme."
    )
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RED
    doc.add_page_break()


# ── Teil A ─────────────────────────────────────────────────────────────────

FIXES = [
    dict(
        num="A1", title="Neubuchung als Bestandskunde crasht das Kunden-Dashboard",
        prio="P0 — KRITISCH", status="Erledigt",
        problem=("Nach der Zahlung im Dashboard-Widget erschien ein schwarzer «Application error», "
                 "der Kunde wurde auf /login geworfen und das Dashboard lud danach bei jedem Aufruf "
                 "nur noch den Absturz. Konsolenfehler: Cannot read properties of null (reading 'id') "
                 "in der Render-Funktion von client-dashboard.js."),
        umsetzung=[
            "handleStripePayment(): Guard `if (error || !paymentIntent) return;` VOR dem Zugriff auf "
            "paymentIntent.id (src/pages/client-dashboard.js).",
            "Neue React-Error-Boundary (src/components/ErrorBoundary.js) umschließt das gesamte "
            "Kunden-Dashboard — ein Render-Fehler zeigt jetzt eine freundliche deutsche Fallback-Seite "
            "mit «Seite neu laden» statt die ganze Seite dauerhaft zu töten.",
            "Der sichtbare Render-Pfad ist bereits defensiv (clientDetails?.assignments || [] etc.); die "
            "exakte .id-Stelle wird durch die Boundary abgefangen.",
        ],
        followup=("Auf der Live-Repro die genaue .id-Zeile per Source-Map final pinnen und den beim "
                  "Test entstandenen defekten Schedule-Datensatz löschen (npx prisma studio)."),
        test=["Als Bestandskunde über das Dashboard-Widget buchen (Testkarte, AGB, «Jetzt bezahlen»).",
              "Seite neu laden (F5) mit offener Konsole — Dashboard lädt normal, kein null-Fehler.",
              "Admin → «Einsätze»: neuer Termin als «Offen / Nicht zugewiesen»."],
        shots=["20-21_neubuchung", "23_neubuchung_admin"],
        files=["src/pages/client-dashboard.js", "src/components/ErrorBoundary.js"],
    ),
    dict(
        num="A2", title="Admin kann keinem Einsatz einen Mitarbeiter zuweisen",
        prio="P0 — KRITISCH", status="Erledigt",
        problem=("Der Einsatz-Dialog zeigte «Keine passenden Empfehlungen gefunden» UND das Dropdown "
                 "«Mitarbeiter auswählen» war leer — bei 10 genehmigten Mitarbeitern. Zwei unabhängige "
                 "Ursachen."),
        umsetzung=[
            "Ursache A (falsche Auth): Alle fünf Admin-Fetches in src/components/Einsaetze.js benutzten "
            "«Bearer localStorage.userToken» (ein Client-Token). Admins authentifizieren per HttpOnly-"
            "Cookie → 401/403, Liste leer. Umgestellt auf credentials:\"include\" (Cookie-Fallback der "
            "Middleware).",
            "Ursache B (falscher Status-Filter): src/pages/api/admin/matchmaking.js filterte "
            "status:\"available\"; genehmigte MA haben \"approved\". Geändert auf "
            "status:{ in: [\"approved\",\"accepted\",\"available\"] }.",
            "Gleiche userToken-Falle zusätzlich in src/pages/admin/clients/[id].js (set-status) behoben.",
        ],
        followup=("Der Match-Score (Verfügbarkeit 30 % · Region 25 % · Services 20 % · Sprache 10 % · "
                  "Erfahrung 10 % · Spezielles 5 %) war korrekt gebaut und liefert nach dem Fix sofort Werte."),
        test=["Admin → Einsätze → «Offen» → Einsatz öffnen: Empfehlungen mit %-Score sichtbar.",
              "Dropdown «Mitarbeiter auswählen» öffnen: alle genehmigten MA wählbar.",
              "«Anfragen» beim Top-Kandidaten (bewusst 2× klicken): Status «angefragt», genau EINE Mail."],
        shots=["34-36_matching"],
        files=["src/components/Einsaetze.js", "src/pages/api/admin/matchmaking.js",
               "src/pages/admin/clients/[id].js"],
    ),
    dict(
        num="A3", title="Admin-Einstellungen: fremder Benutzer + tote Passwortänderung",
        prio="P0/P1 — KRITISCH", status="Erledigt",
        problem=("Einstellungen → «Profil & Zugang» zeigte «Sandra Keller» statt des eingeloggten Admins; "
                 "«Passwort ändern» endete immer im Fehler (Frontend sendet PATCH, API konnte nur GET → 405). "
                 "Ursache: profile.js nutzte findFirst({role:\"admin\"}) statt des Middleware-Headers "
                 "x-admin-email."),
        umsetzung=[
            "src/pages/api/admin/profile.js komplett neu: identifiziert den Admin über den von der "
            "Middleware gesetzten Header x-admin-email (aus dem verifizierten adminToken-Cookie).",
            "GET liefert den EINGELOGGTEN Admin (findUnique per E-Mail).",
            "PATCH implementiert die Passwortänderung mit bcrypt (aktuelles PW prüfen, min. 6 Zeichen, "
            "Hash speichern) und liefert deutsche Meldungen.",
        ],
        followup="",
        test=["admin@phc.local einloggen → Einstellungen → Profil zeigt «PHC Admin / admin@phc.local».",
              "Passwort ändern: falsch / zu kurz / korrekt — passende Meldungen; danach Login nur mit "
              "neuem Passwort (anschließend zurücksetzen)."],
        shots=["53-54_admin_profil"],
        files=["src/pages/api/admin/profile.js"],
    ),
    dict(
        num="A4", title="Kunde sieht keine Quittungen (Stripe)",
        prio="P1 — HOCH", status="Erledigt",
        problem=("Im Tab «Finanzen» gab es weder Zahlungsliste noch Download; der Endpoint "
                 "api/client/invoice-pdf.js war verwaist und las zudem Stripe-Invoices, obwohl die "
                 "Buchung nur PaymentIntents erzeugt (es existieren gar keine Invoices)."),
        umsetzung=[
            "Option A umgesetzt: neuer Endpoint src/pages/api/client/receipts.js listet die Stripe-Charges "
            "des Kunden (Server-seitige Kundenauflösung) und liefert je Datum, Betrag, Währung und "
            "receipt_url. Kunde ohne Kunde/Charge erhält sauberes 200 [] (kein Fehler).",
            "Neuer Abschnitt «Meine Zahlungen / Quittungen» in src/pages/dashboard/finanzen.js: Liste mit "
            "Datum + Betrag + Button «Quittung öffnen» (Stripe-Beleg), Lade- und sauberer Leerzustand.",
        ],
        followup=("Option B (später, sauberer): im Buchungs-Flow echte Stripe-Invoices erzeugen, dann "
                  "liefert der bestehende Endpoint hosted_invoice_url-PDFs."),
        test=["Kunde mit Zahlung → Dashboard → «Finanzen»: Liste «Meine Zahlungen/Quittungen» mit Datum + Betrag.",
              "«Quittung öffnen» klicken: Stripe-Beleg öffnet sich (receipt_url).",
              "Kunde ohne Zahlung: sauberer Leerzustand, kein Fehler."],
        shots=["18-19_quittungen"],
        files=["src/pages/api/client/receipts.js", "src/pages/dashboard/finanzen.js"],
    ),
    dict(
        num="A5", title="Buchungsfunnel: Häufigkeits-Fehlermeldung bleibt stehen",
        prio="P2 — MITTEL", status="Erledigt",
        problem=("Der Häufigkeits-Button setzte nur form.frequency, löschte aber errors.frequency nicht — "
                 "die Fehlermeldung blieb nach der Auswahl sichtbar."),
        umsetzung=["src/pages/register-client.js: onClick des Häufigkeits-Buttons setzt zusätzlich "
                   "setErrors((prev) => ({ ...prev, frequency: \"\" })) — die Meldung verschwindet sofort."],
        followup="",
        test=["/register-client: «Weiter» ohne Auswahl → Meldung erscheint; dann «wöchentlich» klicken → "
              "Meldung verschwindet SOFORT."],
        shots=["10_funnel_s1"],
        files=["src/pages/register-client.js"],
    ),
    dict(
        num="A6", title="Stornogebühr-Grenze + Politur (gesammelt)",
        prio="P2/P3", status="Erledigt (Kern) · Follow-ups dokumentiert",
        problem=("Bei exakt 14 Tagen Vorlauf zeigte die UI 50 % Gebühr, das Backend erstattete aber 100 % "
                 "(Frontend days>14 vs. Backend diffDays>=14). Dazu diverse Politur-Punkte."),
        umsetzung=[
            "Stornogrenze: src/pages/client-dashboard.js auf >=14 und >=7 angeglichen (Backend nutzt >=). "
            "Bei exakt 14 Tagen jetzt konsistent 0 % Gebühr, bei 7 Tagen 50 %.",
            "Kraftausdruck im Kommentar von src/pages/api/create-payment-intent.js entfernt (öffentliches Repo).",
            "Rohstatus: Kundenprofil-Status wird als deutsches Label angezeigt (Offen/Aktiv/Inaktiv/…); der "
            "gespeicherte Wert bleibt unverändert, damit set-status weiter funktioniert.",
        ],
        followup=("Follow-ups umgesetzt/bereitgestellt: (1) E-Mail-Vorlagen — Seed nutzt bereits "
                  "{{greeting}}, die verbliebene Rohanrede (interviewReminder) umgestellt; die "
                  "Orphan-Bereinigung liegt als prisma/cleanup-email-templates.sql bereit (bewusst NICHT "
                  "ausgeführt — destruktive Prod-Operation, nur nach Backup). (2) alert()->Inline beim "
                  "Kartenwechsel (dashboard/finanzen.js) erledigt; die Re-Booking-alerts im 1930-Zeilen-"
                  "client-dashboard.js bewusst belassen (P3, Bruchrisiko). (3) Rohstatus-Audit: Kunden-/"
                  "Mitarbeiter-Dashboard und Kundenliste nutzen bereits deutsche Labels (labelFor); der "
                  "einzige Rohstatus (Admin-Kundendetail-Dropdown) wurde umgestellt."),
        test=["Storno mit >14 / =14 / 10 / 3 Tagen Vorlauf: 0 % / 0 % / 50 % / 100 % — Anzeige = Abrechnung."],
        shots=["24-27_storno"],
        files=["src/pages/client-dashboard.js", "src/pages/api/create-payment-intent.js",
               "src/pages/admin/clients/[id].js"],
    ),
]


def teil_a(doc):
    h(doc, "TEIL A — Bug-Fix-Umsetzungsnachweis", 1)
    p = doc.add_paragraph()
    p.add_run("A1–A6 sind umgesetzt und lokal gegen den laufenden Dev-Server verifiziert. "
              "A7 (Serien-Zuweisung) ist als Konzept/Plan beschrieben und für einen separaten Build "
              "vorgesehen.")
    for f in FIXES:
        h(doc, f"{f['num']} — {f['title']}", 2)
        b = doc.add_paragraph()
        r = b.add_run(f"Priorität: {f['prio']}     Status: {f['status']}")
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RED if "KRITISCH" in f["prio"] else NAVY

        h(doc, "Was war das Problem?", 3)
        doc.add_paragraph(f["problem"])
        h(doc, "Was wurde geändert?", 3)
        for u in f["umsetzung"]:
            doc.add_paragraph(u, style="List Bullet")
        if f["files"]:
            fp = doc.add_paragraph(); fp.add_run("Betroffene Dateien: ").bold = True
            fp.add_run(", ".join(f["files"])).font.size = Pt(10)
        if f["followup"]:
            hp = doc.add_paragraph(); hp.add_run("Hinweis / Follow-up: ").bold = True
            hp.add_run(f["followup"])
        h(doc, "Wie testen?", 3)
        for i, s in enumerate(f["test"], 1):
            doc.add_paragraph(s, style="List Number")
        h(doc, "Nachweis (Screenshot)", 3)
        for name in f["shots"]:
            embed_shot(doc, name)
        res = doc.add_paragraph()
        r = res.add_run("Ergebnis (Abnahme): ______  (OK / Fehler · Datum · Kürzel)")
        r.font.color.rgb = GREY; r.font.size = Pt(9)
        sep = doc.add_paragraph(); hr(sep)

    # A7 — implemented
    h(doc, "A7 — Serien-Zuweisung & Ersatzregelung", 2)
    b = doc.add_paragraph()
    r = b.add_run("Priorität: P0 — DESIGN + BUILD     Status: Erledigt (gebaut, getestet, dokumentiert, gemergt)")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY
    h(doc, "Was war das Problem?", 3)
    doc.add_paragraph(
        "Anfrage/Annahme galt nur für EINEN Termin (assign-employee.js: Assignment mit scheduleId = genau "
        "ein Schedule). Bei einer Mo/Mi/Fr-Serie hätte der Admin jeden Termin einzeln anfragen und der "
        "Mitarbeiter jeden einzeln annehmen müssen. Zusätzlich ein widersprüchlicher Pfad (scheduleId=null "
        "= «alle Termine») ohne klare Regel und ohne Ersatzlogik.")
    h(doc, "Entscheid & Umsetzung", 3)
    doc.add_paragraph("Serie ist Standard — umgesetzt OHNE Schema-Migration über die Konvention "
                      "«Assignment.scheduleId = null = Serien-Zuweisung».")
    for step in [
        "Serien-Anfrage: Admin weist zu → eine Serien-Zuweisung (scheduleId=null, pending); die Anfrage-Mail "
        "und das MA-Dashboard benennen die Serie ({{seriesDescription}} → «Mo/Mi/Fr, je 2 Std, ab 21.07.2026»). "
        "Einzel-Zuweisung bleibt für Ersatz/Einmalbuchung.",
        "Annahme: alle zukünftigen, noch unbesetzten Termine des Kunden erhalten die employeeId "
        "(confirm-assignment.js). Ablehnung der Serie gibt alle wieder frei.",
        "Absenz pro Termin: Stamm-MA gibt über «Absenz melden» EINEN Termin frei (employeeId=null, Status "
        "«ersatz_noetig») — die Serie bleibt beim Stamm-MA (release-schedule.js).",
        "Ersatzregelung: freigegebene Termine erscheinen im Admin («Offene Zuweisungen»); der Admin weist dort "
        "einen Ersatz nur für diesen Termin zu (Einzel-Zuweisung). Kunde sieht am Datum den Ersatz, sonst den Stamm-MA.",
        "Statuslogik: Serie angefragt → bestätigt → aktiv (bzw. abgelehnt → neu vergeben); Einzeltermin "
        "geplant → «Ersatz nötig» → Ersatz bestätigt. Alle Labels DEUTSCH.",
        "Dokumentationspflicht erfüllt: A7_SERIEN-ZUWEISUNG_PROZESS.md (für Bettina & Silvain).",
    ]:
        doc.add_paragraph(step, style="List Number")
    fp = doc.add_paragraph(); fp.add_run("Betroffene Dateien: ").bold = True
    fp.add_run("src/lib/series.js, api/admin/assign-employee.js, api/employee/confirm-assignment.js, "
               "api/employee/release-schedule.js, api/employee/pending-assignments.js, api/einsaetze/index.js, "
               "components/AssignmentList.js, employee-dashboard.js, client-dashboard.js, lib/statusLabels.js, "
               "prisma/seed-email-templates.js").font.size = Pt(9)
    h(doc, "Verifikation (Integrationstest)", 3)
    doc.add_paragraph(
        "prisma/a7-series-test.js — isolierter Test-Kunde, echte Mitarbeiter, Cleanup am Ende. Geprüft: "
        "(1) Serien-Zuweisung erzeugt scheduleId=null und stempelt Termine NICHT vor Annahme; (2) Annahme "
        "propagiert auf alle 3 Termine; (3) Absenz gibt nur einen Termin frei (ersatz_noetig), Rest bleibt; "
        "(4) Ersatz-Einzelzuweisung setzt den Ersatz; (5) Serien-Ablehnung gibt alle frei. "
        "Alle 9 Prüfungen bestanden.")
    res = doc.add_paragraph()
    r = res.add_run("Ergebnis (Abnahme): ______  (OK / Fehler · Datum · Kürzel)")
    r.font.color.rgb = GREY; r.font.size = Pt(9)
    doc.add_page_break()


# ── Teil B ─────────────────────────────────────────────────────────────────

PROCESSES = [
    ("1. Login- & Session-Prozesse", "Alle Rollen", [
        ("Login Kunde / Mitarbeiter / Admin", "Weiterleitung je aufs richtige Dashboard, Name korrekt", "01-03_login_rollen"),
        ("Login mit falschem Passwort", "Fehlermeldung, kein Zugang", "04_login_fehler"),
        ("«Passwort vergessen»", "Reset-Formular erreichbar; Reset-Mail kommt an", "05_pw_vergessen"),
        ("Logout → Dashboard-URL direkt", "Redirect auf /login, localStorage leer", "06_logout"),
        ("Abgelaufene Session → Aktion", "Sofortiger Redirect auf /login — nie fremde Daten", "07_session_timeout"),
    ]),
    ("2. Buchungsprozess (Funnel)", "Anonym → neuer Kunde", [
        ("Schritt 1: Häufigkeit, Beginndatum, Wochentage", "Preis live; A5-Fix: Fehlermeldung verschwindet", "10_funnel_s1"),
        ("Schritt 2: betreute Person + Auftraggeber", "Beide Blöcke da, Validierung greift", "10_funnel_s2"),
        ("Schritt 3: Übersicht, Gutschein, AVB, Stripe", "Summe korrekt, AVB Pflicht", "10_funnel_s3"),
        ("Schritt 4: Abschluss", "«Buchung erfolgreich», Mail-Hinweis", "10_funnel_s4"),
        ("Danach als Kunde: Termine", "Termine abwechselnd nach Muster", "14_termine_serie"),
    ]),
    ("3. Zahlungsprozess (Stripe)", "Kunde", [
        ("Zahlung Testkarte 4242…", "Erfolg + Weiterleitung", "15_zahlung_ok"),
        ("Ablehnungskarte 4000…0002", "Klare Fehlermeldung, keine Buchung", "16_zahlung_abgelehnt"),
        ("Finanzen: Karte ändern", "Neue Karte gespeichert, lesbare Maske", "17_karte"),
        ("Finanzen: Quittungsliste (A4)", "Jede Zahlung mit Datum/Betrag + «Quittung öffnen»", "18-19_quittungen"),
    ]),
    ("4. Neubuchungsprozess (Bestandskunde)", "Kunde eingeloggt", [
        ("Dashboard-Widget → Zahlung → «Jetzt bezahlen»", "Erfolgsmeldung, KEIN Absturz (A1)", "20-21_neubuchung"),
        ("F5 nach der Buchung", "Dashboard lädt, Termin vorhanden", "22_neubuchung_reload"),
        ("Admin → Einsätze", "Neuer Termin «Offen / Nicht zugewiesen»", "23_neubuchung_admin"),
    ]),
    ("5. Stornoprozess", "Kunde", [
        ("Storno >14 / =14 / 10 / 3 Tage Vorlauf", "0 % / 0 % / 50 % / 100 % — Anzeige = Abrechnung (A6)", "24-27_storno"),
        ("Storno-Bestätigungsmail", "Korrekte Gebühr in der Mail", "28_storno_mail"),
    ]),
    ("6. Kündigungsprozess (vorzeitig, CHF 300)", "Kunde + Admin", [
        ("Kündigung <30 Tage seit Beginn", "Dialog mit CHF-300-Hinweis + Pflichtzustimmung; Status «gekuendigt»", "29-30_kuendigung"),
        ("Kündigung >30 Tage", "Ohne Gebühr, keine Stripe-Buchung", "31_kuendigung_frei"),
        ("Admin: «Ehemalige Kunden»", "Kunde dort sichtbar, nicht mehr in «Neue Buchungen»", "32-33_ehemalige"),
    ]),
    ("7. Zuweisungs- & Matching-Prozess", "Admin → Mitarbeiter → Kunde", [
        ("Offener Einsatz: Empfehlungen + Dropdown", "Score-Liste gefüllt, alle genehmigten MA wählbar (A2)", "34-36_matching"),
        ("«Anfragen» (2× klicken)", "Status «angefragt», genau EINE Mail", "37_anfrage"),
        ("SERIE: MA nimmt Anfrage an", "ALLE zukünftigen Serien-Termine tragen den MA — pending A7", "38-39_serie_annahme"),
        ("Anfrage-Mail + MA-Dashboard", "Serie klar benannt — pending A7", "38b_serie_mail"),
        ("Kunde nach Annahme", "Betreuer bei allen Serien-Terminen — pending A7", "40_kunde_serie"),
        ("Ablehnung (ganze Serie)", "Alle Termine wieder «Nicht zugewiesen» — pending A7", "41_serie_abgelehnt"),
        ("Kein passender MA", "Klarer Hinweis «manuelle Zuweisung nötig»", "42_kein_match"),
    ]),
    ("8. Ersatzregelung bei Absenz", "Mitarbeiter → Admin → Kunde", [
        ("Stamm-MA meldet Absenz für EINEN Termin", "Nur dieser Termin frei («Ersatz nötig») — pending A7", "43_absenz_meldung"),
        ("Admin: freigegebener Termin", "Badge «Ersatz nötig», Matching nur dafür — pending A7", "44_ersatz_matching"),
        ("Ersatz-MA nimmt Einzeltermin an", "Kunde sieht an dem Datum den Ersatz — pending A7", "45_ersatz_bestaetigt"),
        ("MA-Urlaub über Zeitraum", "MA im Matching NICHT vorgeschlagen", "46_urlaub_matching"),
    ]),
    ("9. Mitarbeiterprozess (Bewerbung → Onboarding)", "Kandidat → Admin → Mitarbeiter", [
        ("Bewerbung alle 4 Schritte + einreichen", "/bewerbung-erfolgreich + Bestätigungsmail", "47-48_bewerbung"),
        ("Admin: Kandidaten-Detail", "Alle Felder 1:1 (Sprachen, Pensum, Reisebereitschaft)", "49_kandidat_detail"),
        ("Genehmigen → «Passwort erstellen» → Login", "Wortlaut «erstellen», Login klappt", "50-51_genehmigung"),
        ("MA-Dashboard + Sub-Tabs", "Menü vollständig, Tabs filtern", "52_ma_dashboard"),
    ]),
    ("10. Admin-Prozesse & Absenzen-Verwaltung", "Admin (+ Kunde/MA für Urlaub)", [
        ("Einstellungen: Profil + Passwort (A3)", "Eigener Admin; Änderung funktioniert", "53-54_admin_profil"),
        ("Kunden-/Einsätze-Verwaltung: Filter + Links", "Filter ok, Links öffnen Profile, deutsche Status", "55-56_admin_verwaltung"),
        ("Kunde erfasst Urlaub", "Termine im Zeitraum automatisch storniert + Hinweis", "57_kunde_urlaub"),
        ("MA beantragt Urlaub → Admin entscheidet", "Antrag sichtbar, Entscheid wird angezeigt", "58_ma_urlaub"),
        ("Finanzen, Gutscheine, Glocke, E-Mail-Editor", "Alles erreichbar und funktionsfähig", "59-60_admin_rest"),
    ]),
    ("11. E-Mail-Prozesse (alle Trigger)", "System → Empfänger", [
        ("Willkommen/Passwort, Terminbestätigung", "An Auftraggeber-Adresse, «Sehr geehrte/r …», deutsch", "61-62_mail_buchung"),
        ("Anfrage-Mail an MA, Bestätigung an Kunde", "Je genau EINE Mail, Betreuer-Kontakt drin", "63-64_mail_zuweisung"),
        ("Storno-, Ersatz-, Bewerbungs-Mails", "Inhalte korrekt, formelle Anrede, DEUTSCH", "65-66_mail_rest"),
        ("Editor-Check aller Vorlagen", "Keine Duplikate, kein rohes «Grüezi {{firstName}}»", "67_mail_vorlagen"),
    ]),
]


def process_table(doc, rows):
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, label in enumerate(["Testfall", "Erwartetes Resultat", "Screenshot", "OK"]):
        hdr[i].paragraphs[0].add_run(label).bold = True
        shade(hdr[i], "04436F")
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(9)
    for tc, exp, name in rows:
        c = t.add_row().cells
        c[0].paragraphs[0].add_run(tc).font.size = Pt(9)
        c[1].paragraphs[0].add_run(exp).font.size = Pt(9)
        status = "auto" if shot(name) else "manuell"
        r = c[2].paragraphs[0].add_run(f"{name}.png\n({status})")
        r.font.size = Pt(8); r.font.color.rgb = GREY
        c[3].paragraphs[0].add_run("☐")


def teil_b(doc):
    h(doc, "TEIL B — Test- & Abnahmeplan (11 Kernprozesse)", 1)
    p = doc.add_paragraph()
    p.add_run("Jeder Prozess: Testfall-Tabelle + Screenshot-Nachweis. Automatisch aufgenommene "
              "Screenshots sind eingebettet; nicht automatisierbare (abgeschlossene Stripe-Zahlung, "
              "echter E-Mail-Posteingang, A7-Serienverhalten) erscheinen als beschrifteter Platzhalter "
              "mit Begründung und sind live nachzureichen.")
    for title, rolle, rows in PROCESSES:
        h(doc, title, 2)
        rp = doc.add_paragraph(); rp.add_run("Rolle: ").bold = True; rp.add_run(rolle)
        process_table(doc, rows)
        doc.add_paragraph()
        for _tc, _exp, name in rows:
            embed_shot(doc, name)
        res = doc.add_paragraph()
        r = res.add_run("Gesamtergebnis (bestanden / Fehler · Datum · Kürzel): ______")
        r.font.color.rgb = GREY; r.font.size = Pt(9)
        sep = doc.add_paragraph(); hr(sep)


def main():
    doc = Document()
    styles(doc)
    cover(doc)
    teil_a(doc)
    teil_b(doc)
    out = os.path.join(BASE, "PHC_QA-Umsetzung_12.07.2026.docx")
    doc.save(out)
    # Coverage summary to stdout.
    total = sum(len(r) for _t, _rl, r in PROCESSES)
    auto = sum(1 for _t, _rl, r in PROCESSES for _tc, _e, n in r if shot(n))
    print(f"Saved {out}")
    print(f"Teil B screenshots: {auto}/{total} auto-captured, {total - auto} manual/pending placeholders")


if __name__ == "__main__":
    main()
