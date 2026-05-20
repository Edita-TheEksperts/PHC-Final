"""
Builds PHC_Test_Dokumentation_FULL.docx — a comprehensive companion to the
basic doc. Adds three things in front of the standard 48-item audit content:

  1. Test-Zugangsdaten (credentials table for QA)
  2. Buchungs-Funnel walkthrough (Register-Client step by step)
  3. Bewerbungs-Funnel walkthrough (Employee-Register step by step)

After those, the same `build()` function from build_test_doc.py is invoked
to append the full audit content. So the FULL doc = walkthroughs + everything
the basic doc has, in one self-contained file.

Screenshots:
  REG-Client-1..4   captured per step via ?step=N URL parameter
  REG-Employee      stacked "all steps" capture
  REG-Employee-1..4 individual step captures

Run:  python build_full_doc.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Re-use everything from the basic builder: helpers, audit content, the lot.
from build_test_doc import (
    build as build_audit_content,
    apply_doc_styles,
    add_title_and_vorwort,
    SCREENSHOT_DIR,
    find_screenshots,
    screenshot_placeholder,
    add_horizontal_line,
)


def shot_or_placeholder(doc, name, caption):
    """Embed screenshots/<name>.png if it exists, otherwise a placeholder box."""
    path = os.path.join(SCREENSHOT_DIR, f"{name}.png")
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(6.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    else:
        screenshot_placeholder(doc, caption)


def add_credentials_section(doc):
    doc.add_heading("Test-Zugangsdaten", level=1)
    doc.add_paragraph(
        "Diese drei Konten sind im Test-System angelegt und sollten für die "
        "Abnahme verwendet werden. Die Konten werden idempotent durch "
        "`npm run seed:test-users` erzeugt — re-running überschreibt nichts."
    )
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Rolle"
    hdr[1].text = "E-Mail"
    hdr[2].text = "Passwort"
    hdr[3].text = "Angezeigter Name"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs: r.bold = True

    rows = [
        ("Admin",     "admin@phc.local",                "TestAdmin2026!",   "PHC Admin"),
        ("Kunde",     "edita.latifi@the-eksperts.com",  "TestEdita2026!",   "Edita Latifi"),
        ("Mitarbeiter","fisnik.test@phc.local",         "TestFisnik2026!",  "Fisnik Salihu"),
        ("Kandidat",  "lina.kandidat@phc.local",        "TestFisnik2026!",  "Lina Kandidat (pending)"),
    ]
    for r in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v

    p = doc.add_paragraph()
    p.add_run(
        "Hinweise: der reale Outlook-Account "
        "fisnik.salihu@the-eksperts.com existiert bereits als Admin im "
        "Live-System und kann daher nicht parallel als Mitarbeiter-Login "
        "dienen — deshalb das Test-Suffix .test@phc.local. Edita's reale "
        "Mailadresse wird bewusst als Kundenkonto verwendet, damit sich der "
        "E-Mail-Versand (F-12) im echten Outlook-Postfach visuell verifizieren "
        "lässt."
    ).italic = True
    doc.add_page_break()


def add_buchungs_funnel(doc):
    doc.add_heading("Teil A — Buchungs-Funnel (Kunde)", level=1)
    doc.add_paragraph(
        "Der Buchungs-Funnel führt einen Kunden in vier Schritten durch die "
        "Erstbuchung. Diese Sektion zeigt jeden Schritt einzeln, damit ihr im "
        "Test prüfen könnt, ob alle relevanten Felder vorhanden sind und ob die "
        "Übergänge zwischen den Schritten funktionieren."
    )
    doc.add_paragraph(
        "Start: /betreuung-zuhause-organisieren (oder /register-client) als "
        "anonymer Besucher. Zum schnellen Springen zwischen Schritten kann der "
        "URL-Parameter ?step=N verwendet werden."
    )

    steps = [
        {
            "n": 1,
            "title": "Schritt 1 — Dienstleistungen & Termine",
            "what": (
                "Der Kunde wählt eine Haupt-Dienstleistung (Haushaltshilfe, "
                "Freizeit, Einkäufe, Terminbegleitung) und legt für jeden "
                "gewünschten Wochentag fest, welche Subservices an diesem Tag "
                "erbracht werden sollen. Die Auswahl ist tagesabhängig — "
                "Montag kann z. B. Kochen sein, Mittwoch Spazieren (F-01). "
                "Stundenanzahl und Preis werden live berechnet (F-02)."
            ),
            "fields": [
                "Haupt-Dienstleistung (Pflichtfeld)",
                "Subservices pro Wochentag-Zeile (mehrfach wählbar)",
                "Wochentag, Uhrzeit, Stundenanzahl pro Zeile",
                "Beginndatum (frühestens 10 Tage in der Zukunft)",
                "Frequenz (einmalig / wöchentlich / monatlich)",
            ],
            "expected": (
                "Mehrere Schedule-Zeilen sind möglich. Subservices pro Zeile "
                "sind unabhängig wählbar. Preis aktualisiert sich automatisch "
                "(CHF 59 × Stunden)."
            ),
            "shot": "REG-Client-1",
        },
        {
            "n": 2,
            "title": "Schritt 2 — Angaben zur betreuten Person",
            "what": (
                "Der Kunde gibt persönliche Daten der betreuten Person und "
                "(separat) die Daten der anfragenden Person (Auftraggeber) "
                "ein. F-08: die Auftraggeber-Adresse wird seit dem Audit "
                "explizit getrennt erfasst — sie kann von der Adresse der "
                "betreuten Person abweichen."
            ),
            "fields": [
                "Anrede, Vor- und Nachname (betreute Person)",
                "Strasse + Hausnummer (F-03 / F-28)",
                "PLZ, Ort, Kanton",
                "Geburtsdatum, Telefon",
                "Separate Auftraggeber-Adresse: Strasse, Hausnummer, PLZ, Ort, Kanton, E-Mail (F-08)",
            ],
            "expected": (
                "Beide Adressblöcke sind sichtbar, alle Pflichtfelder klar "
                "gekennzeichnet, Validation greift bei leeren Pflichtfeldern."
            ),
            "shot": "REG-Client-2",
        },
        {
            "n": 3,
            "title": "Schritt 3 — Zahlung",
            "what": (
                "Übersicht aller gebuchten Termine, Stundensumme und "
                "Gesamtpreis. Hier wird das Stripe-Zahlungselement angezeigt. "
                "Optional kann ein Gutscheincode eingegeben werden (F-30). "
                "Vor der Zahlung müssen AVB akzeptiert werden."
            ),
            "fields": [
                "Termin-Übersichtsblock (Datum, Uhrzeit, Stunden, Subservice)",
                "Gesamtpreis-Anzeige (CHF 59 × Stunden)",
                "Gutscheincode-Eingabe (optional)",
                "Stripe-Zahlungselement (Test-Karte: 4242 4242 4242 4242)",
                "AVB-Checkbox + Datenschutz-Hinweis",
            ],
            "expected": (
                "Stripe-Element lädt, Zahlung mit Test-Karte ist erfolgreich, "
                "Weiterleitung zu Schritt 4."
            ),
            "shot": "REG-Client-3",
        },
        {
            "n": 4,
            "title": "Schritt 4 — Bestätigung",
            "what": (
                "Erfolgsseite nach abgeschlossener Zahlung. Hier wird "
                "bestätigt, dass die Buchung registriert wurde und eine "
                "Willkommens-E-Mail an die Auftraggeber-Adresse versendet "
                "wurde (F-12)."
            ),
            "fields": [
                "Bestätigungs-Text (\"Buchung erfolgreich\")",
                "Hinweis: \"Eine E-Mail wurde an … gesendet\"",
                "Link/Button zum Login",
            ],
            "expected": (
                "Willkommens-E-Mail landet im Postfach der Auftraggeber-Adresse "
                "(im Test-Setup: edita.latifi@the-eksperts.com). E-Mail enthält "
                "Passwort-erstellen-Link (F-16)."
            ),
            "shot": "REG-Client-4",
        },
    ]

    for s in steps:
        doc.add_heading(s["title"], level=2)
        doc.add_heading("Was passiert hier?", level=3)
        doc.add_paragraph(s["what"])
        doc.add_heading("Felder in diesem Schritt", level=3)
        for f in s["fields"]:
            doc.add_paragraph(f, style="List Bullet")
        doc.add_heading("Screenshot", level=3)
        shot_or_placeholder(doc, s["shot"], f"{s['title']} — automatisch aufgenommen")
        doc.add_heading("Erwartetes Ergebnis", level=3)
        doc.add_paragraph(s["expected"])
        sep = doc.add_paragraph()
        add_horizontal_line(sep)

    doc.add_page_break()


def add_bewerbungs_funnel(doc):
    doc.add_heading("Teil B — Bewerbungs-Funnel (Mitarbeiter / Kandidat)", level=1)
    doc.add_paragraph(
        "Der Bewerbungs-Funnel führt Bewerber:innen in vier Schritten durch "
        "die Erstbewerbung. Im Live-System sind die Schritte mit "
        "{display: hidden} entkoppelt — für die Test-Doku haben wir alle "
        "vier Schritte sichtbar gemacht, damit jedes Feld dokumentiert ist."
    )

    steps = [
        {
            "n": 1,
            "title": "Schritt 1 — Persönliche Informationen",
            "what": (
                "Grunddaten der bewerbenden Person. Diese Daten landen im "
                "Employee-Datensatz und sind später im Admin-Detail sichtbar "
                "(F-44 / F-45)."
            ),
            "fields": [
                "Anrede (Herr / Frau / Divers)",
                "Vor- und Nachname",
                "E-Mail, Telefon",
                "Strasse, Hausnummer, PLZ, Ort, Kanton",
                "Geburtsdatum",
                "Nationalität, Aufenthaltsbewilligung",
            ],
            "expected": "Alle Pflichtfelder validiert. Weiter zu Schritt 2.",
            "shot": "REG-Employee-1",
        },
        {
            "n": 2,
            "title": "Schritt 2 — Weitere Informationen",
            "what": (
                "Erfahrungs-, Sprach- und Kommunikationsangaben. F-34: alle "
                "Sprachen, Kommunikationsmerkmale und Ernährungserfahrungen "
                "müssen 1:1 in der Datenbank landen und im Admin-Detail "
                "sichtbar sein."
            ),
            "fields": [
                "Erfahrungsjahre",
                "Sprachen (Mehrfachauswahl: DE / FR / IT / EN …)",
                "Andere Sprache (Freitext)",
                "Kommunikationsmerkmale (z. B. geduldig, empathisch)",
                "Ernährungserfahrung (z. B. vegetarisch, diabetikergerecht)",
                "Spezielle Trainings (Demenz, Palliative, Wundpflege …)",
            ],
            "expected": "Mehrfachauswahl funktioniert für alle drei Listen.",
            "shot": "REG-Employee-2",
        },
        {
            "n": 3,
            "title": "Schritt 3 — Arbeitsbereitschaft",
            "what": (
                "Verfügbarkeit, gewünschtes Arbeitspensum (F-05 "
                "Mehrfachauswahl), angebotene Dienstleistungen (F-35), "
                "Reisebereitschaft (F-45) und persönliche Eigenschaften "
                "(Raucher, Haustiere, Führerschein)."
            ),
            "fields": [
                "Verfügbar ab (Datum)",
                "Verfügbare Wochentage (Mehrfachauswahl)",
                "Arbeitspensum (40h, 32h, 24h, 16h, 8h — mehrfach wählbar, F-05)",
                "Angebotene Dienstleistungen (Haushaltshilfe, Freizeit, Einkäufe, Terminbegleitung)",
                "Wie weit reisebereit (km, F-45)",
                "Wochenend-, Nachtschicht-, On-Call-Bereitschaft",
                "Raucher / Nichtraucher, Haustier-Bereitschaft, Führerschein",
            ],
            "expected": (
                "Pensum-Auswahl ist eine Mehrfachauswahl (Buttons werden farbig "
                "wenn ausgewählt). Dienstleistungs-Liste verwendet dieselben "
                "Bezeichnungen wie die Kundenseite (F-35 Kongruenz)."
            ),
            "shot": "REG-Employee-3",
        },
        {
            "n": 4,
            "title": "Schritt 4 — Abschluss",
            "what": (
                "Dokument-Upload (CV, Pass, Diplom, Führerschein), "
                "Bestätigung der AGB und Absenden der Bewerbung. Nach "
                "Submit Weiterleitung auf /bewerbung-erfolgreich (F-06)."
            ),
            "fields": [
                "CV (Pflicht-Upload)",
                "Pass / Aufenthaltsbewilligung (Pflicht)",
                "Diplom / Zertifikate (optional)",
                "Führerschein (optional, wenn hasLicense=true)",
                "AGB-Akzeptanz-Checkbox",
                "Submit-Button",
            ],
            "expected": (
                "Submit erfolgreich → Weiterleitung auf /bewerbung-erfolgreich. "
                "Bewerbungs-E-Mail (applicantConfirmation Template) landet im "
                "Postfach des Bewerbers."
            ),
            "shot": "REG-Employee-4",
        },
    ]

    for s in steps:
        doc.add_heading(s["title"], level=2)
        doc.add_heading("Was passiert hier?", level=3)
        doc.add_paragraph(s["what"])
        doc.add_heading("Felder in diesem Schritt", level=3)
        for f in s["fields"]:
            doc.add_paragraph(f, style="List Bullet")
        doc.add_heading("Screenshot", level=3)
        shot_or_placeholder(doc, s["shot"], f"{s['title']} — automatisch aufgenommen")
        doc.add_heading("Erwartetes Ergebnis", level=3)
        doc.add_paragraph(s["expected"])
        sep = doc.add_paragraph()
        add_horizontal_line(sep)

    # Bonus: the full stacked version (all 4 steps in one giant screenshot)
    doc.add_heading("Alle Schritte gestapelt (Übersicht)", level=2)
    doc.add_paragraph(
        "Zur visuellen Übersicht hier alle vier Schritte als ein zusammenhängendes "
        "Bild — zeigt sofort, dass im Bewerbungsformular kein einziges Feld "
        "ausgelassen wurde."
    )
    shot_or_placeholder(doc, "REG-Employee", "Bewerbungsformular — alle 4 Schritte gestapelt")
    doc.add_page_break()


def build_full(doc):
    """Build the comprehensive FULL doc: title → creds → walkthroughs → audit."""
    apply_doc_styles(doc)
    add_title_and_vorwort(doc, subtitle="Vollständige Test- & Walkthrough-Dokumentation")
    add_credentials_section(doc)
    add_buchungs_funnel(doc)
    add_bewerbungs_funnel(doc)
    # Audit content WITHOUT its own title (we already have one).
    build_audit_content(doc, include_title=False)


if __name__ == "__main__":
    doc = Document()
    build_full(doc)
    out_path = "PHC_Test_Dokumentation_FULL.docx"
    try:
        doc.save(out_path)
        print(f"Saved: {out_path}")
    except PermissionError:
        out_path = "PHC_Test_Dokumentation_FULL_v2.docx"
        doc.save(out_path)
        print(f"Saved: {out_path}  (original was locked — close it in Word and re-run to overwrite)")
